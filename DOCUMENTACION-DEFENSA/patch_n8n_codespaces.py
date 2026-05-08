"""Patch puntual: cambia párrafo TypeScript y añade n8n placeholder en Fase 5."""
import copy
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r'C:\Users\shair\Downloads\Memoria_v2.0_Sofia.docx'
doc = Document(OUT)
paras = doc.paragraphs

def set_text(p, txt):
    rpr = None
    if p.runs and p.runs[0]._r.find(qn('w:rPr')) is not None:
        rpr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr')))
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    nr = p.add_run(txt)
    if rpr is not None:
        nr._r.insert(0, rpr)

def insert_placeholder_after(para, label):
    """Inserta un párrafo de placeholder [CAPTURA: ...] después del párrafo dado."""
    # Nuevo párrafo con estilo igual al existente
    new_p = OxmlElement('w:p')
    # Copiar propiedades de párrafo del original
    orig_pPr = para._p.find(qn('w:pPr'))
    if orig_pPr is not None:
        new_p.append(copy.deepcopy(orig_pPr))
    # Run con texto placeholder
    r = OxmlElement('w:r')
    t = OxmlElement('w:t'); t.text = f'(Insertar aquí captura: {label})'
    r.append(t)
    new_p.append(r)
    para._p.addnext(new_p)
    return new_p

# ── 1. Cambiar párrafo [264] — TypeScript → Codespaces + Grafana + Prometheus ──
p264 = paras[264]
if 'TypeScript ha sido el mayor' in p264.text:
    set_text(p264,
        'En cuanto a herramientas, GitHub Codespaces ha sido una revelación: poder trabajar en un '
        'entorno completamente configurado desde el navegador, sin instalar nada en local, y que '
        'ese mismo entorno funcione idéntico para la demo ante el tribunal, simplifica muchísimo '
        'la logística. Grafana y Prometheus también han superado mis expectativas: al principio '
        'pensaba que serían herramientas difíciles de integrar, pero con prom-client y el endpoint '
        '/metrics del backend, en pocas horas ya tenía dashboards con datos reales del sistema. '
        'Ver el pico de intentos de login en tiempo real durante el ataque de fuerza bruta, '
        'reflejado automáticamente en la gráfica de Grafana, es exactamente el tipo de '
        'visualización que hace el proyecto memorable para el tribunal.'
    )
    print('✓ Párrafo TypeScript reemplazado por Codespaces+Grafana+Prometheus')

# ── 2. Añadir placeholder de n8n en Fase 5 (después del párrafo [227]) ───────
p227 = paras[227]  # 'Grafana fue lo más visual...'
# Insertar placeholder de captura n8n después
insert_placeholder_after(p227,
    'Flujo de automatización n8n — nodos Webhook → Set → Send Email para alertas SOS'
)
print('✓ Placeholder n8n insertado en Fase 5')

# Insertar también texto explicativo de n8n antes del placeholder
new_txt = OxmlElement('w:p')
orig_pPr = p227._p.find(qn('w:pPr'))
if orig_pPr is not None:
    new_txt.append(copy.deepcopy(orig_pPr))
r = OxmlElement('w:r')
t = OxmlElement('w:t')
t.text = ('n8n fue la pieza que unió todo el sistema de alertas. Configuré un workflow con tres '
          'nodos: un Webhook que recibe la petición POST del backend cuando se detecta un ataque '
          'crítico o se pulsa el botón SOS, un nodo Set que formatea el mensaje con el nombre '
          'del cliente y el nivel de urgencia, y un nodo Send Email que envía la alerta al SOC. '
          'La clave fue que la URL del webhook tiene que ser la del dominio público '
          '(serveo.net), no localhost, porque n8n corre dentro de Docker y no puede resolver '
          'peticiones a sí mismo por localhost.')
r.append(t)
new_txt.append(r)
p227._p.addnext(new_txt)
print('✓ Texto explicativo n8n añadido')

doc.save(OUT)
print(f'OK — Guardado: {OUT}')
