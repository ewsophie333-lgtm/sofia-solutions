"""
Final DOCX Refinement — Sinergy of Natural Tone and Technical Accuracy.
"""
import copy
from docx import Document
from docx.oxml.ns import qn

OUT = r'C:\Users\shair\Downloads\Memoria_v2.0_Sofia.docx'
doc = Document(OUT)
paras = doc.paragraphs

def set_text(p, txt):
    rpr = None
    if p.runs and p.runs[0]._r.find(qn('w:rPr')) is not None:
        rpr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr')))
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    nr = p.add_run(txt)
    if rpr is not None:
        nr._r.insert(0, rpr)

# Simplificamos las frases iniciales de las conclusiones
p262 = paras[262]
set_text(p262, 
    'Poder ver en tiempo real lo que pasa en el sistema gracias a las gráficas de Grafana y Prometheus ha sido clave. '
    'Ya no son solo datos aburridos que nadie lee, sino que ves el ataque ocurriendo en directo y puedes reaccionar. '
    'Además, n8n hace que el sistema responda por sí solo: si alguien intenta entrar a la fuerza o si el cliente '
    'pulsa el botón de emergencia SOS, llega un correo de alerta al SOC de forma inmediata.'
)

p263 = paras[263]
set_text(p263,
    'El uso de GitHub Codespaces ha sido un gran acierto, porque me ha permitido olvidarme de configurar el PC. '
    'Simplemente abres el navegador y ya tienes todo listo para funcionar. Gracias a que usamos Docker, el proyecto '
    'se puede desplegar en segundos en cualquier ordenador y sé que funcionará exactamente igual que en mi equipo, '
    'lo que da mucha tranquilidad para el día de la defensa.'
)

doc.save(OUT)
print(f'OK — Conclusiones simplificadas en {OUT}')
