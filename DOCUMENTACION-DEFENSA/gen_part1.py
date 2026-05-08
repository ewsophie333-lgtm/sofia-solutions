"""
Memoria TFG Sofia Solutions v2.0 — Generador completo
Reescribe TODAS las secciones basándose en el contenido definitivo.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, shutil

INPUT  = r'C:\Users\shair\Downloads\Memoria. Entrega 4.docx'
OUTPUT = r'C:\Users\shair\Desktop\sofia-solutions\sofia-solutions\Memoria_v2.0_Sofia.docx'

# ── helpers ────────────────────────────────────────────────────────────────

def body(doc, txt, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(18)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(txt)
    r.font.name = 'Verdana'; r.font.size = Pt(12)
    return p

def h1(doc, txt):
    p = doc.add_heading(txt, level=1)
    for r in p.runs:
        r.font.name = 'Verdana'; r.font.size = Pt(14); r.bold = True
    return p

def h2(doc, txt):
    p = doc.add_heading(txt, level=2)
    for r in p.runs:
        r.font.name = 'Verdana'; r.font.size = Pt(12); r.bold = True
    return p

def bullet(doc, txt):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('• ' + txt)
    r.font.name = 'Verdana'; r.font.size = Pt(12)
    return p

def code(doc, filename, snippet):
    t = doc.add_table(rows=2, cols=1)
    t.style = 'Normal Table'
    hdr = t.rows[0].cells[0]
    hdr.paragraphs[0].clear()
    hr = hdr.paragraphs[0].add_run(filename)
    hr.bold = True; hr.font.name = 'Verdana'; hr.font.size = Pt(10)
    hdr.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'D9E1F2')
    hdr._tc.get_or_add_tcPr().append(shd)
    cc = t.rows[1].cells[0]
    cc.paragraphs[0].clear()
    cr = cc.paragraphs[0].add_run(snippet)
    cr.font.name = 'Courier New'; cr.font.size = Pt(9)
    doc.add_paragraph()

def tbl_header(t, *cols):
    row = t.add_row()
    for i, c in enumerate(cols):
        cell = row.cells[i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(c)
        r.bold = True; r.font.name = 'Verdana'; r.font.size = Pt(10)
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '1F3564')
        cell._tc.get_or_add_tcPr().append(shd)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

def tbl_row(t, *cols):
    row = t.add_row()
    for i, c in enumerate(cols):
        cell = row.cells[i]
        cell.paragraphs[0].clear()
        r = cell.paragraphs[0].add_run(str(c))
        r.font.name = 'Verdana'; r.font.size = Pt(10)

def mk_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Normal Table'
    tbl_header(t, *headers)
    for row in rows:
        tbl_row(t, *row)
    doc.add_paragraph()
    return t

def pb(doc):
    doc.add_page_break()

def caption(doc, n, desc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'Figura {n}: {desc}')
    r.italic = True; r.font.name = 'Verdana'; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x44,0x44,0x44)

def placeholder(doc, desc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'[CAPTURA: {desc}]')
    r.font.name = 'Verdana'; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x77,0x77,0x77)

# ── abrir doc base ─────────────────────────────────────────────────────────
doc = Document(INPUT)

# Eliminar TODOS los párrafos existentes manteniendo el documento
from docx.oxml.ns import qn as ns_qn
body_el = doc.element.body
for child in list(body_el):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag in ('p', 'tbl', 'sectPr'):
        if tag != 'sectPr':
            body_el.remove(child)

# Márgenes
sec = doc.sections[0]
sec.left_margin=Cm(3); sec.right_margin=Cm(2)
sec.top_margin=Cm(2); sec.bottom_margin=Cm(2)

# Encabezado
sec.header.paragraphs[0].clear()
sec.header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr=sec.header.paragraphs[0].add_run('Sofia Solutions — Laura Sofia Gomez Salazar')
hr.font.name='Verdana'; hr.font.size=Pt(9)

# Pie de página
sec.footer.paragraphs[0].clear()
sec.footer.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
fld1=OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'),'begin')
ins=OxmlElement('w:instrText'); ins.text='PAGE'
fld2=OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'),'end')
rr=sec.footer.paragraphs[0].add_run()
rr._r.append(fld1); rr._r.append(ins); rr._r.append(fld2)

# ── PORTADA ────────────────────────────────────────────────────────────────
p=doc.add_paragraph()
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('TRABAJO DE FIN DE GRADO\nCiclo Superior de Administración de Sistemas Informáticos en Xarxa')
r.font.name='Verdana'; r.font.size=Pt(14); r.bold=True
doc.add_paragraph()
p=doc.add_paragraph()
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('SOFIA SOLUTIONS')
r.font.name='Verdana'; r.font.size=Pt(28); r.bold=True
r.font.color.rgb=RGBColor(0x1F,0x35,0x64)
doc.add_paragraph()
p=doc.add_paragraph()
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Plataforma de Ciberseguridad Gestionada para PYMEs\nSOC, Monitorización y Auditoría de Vulnerabilidades')
r.font.name='Verdana'; r.font.size=Pt(14)
doc.add_paragraph()
placeholder(doc, 'Logo Sofia Solutions')
doc.add_paragraph()
info = [
    ('Autora:', 'Laura Sofía Gómez Salazar'),
    ('NIA:', '13225255'),
    ('Tutor:', 'Isaac García Simarro'),
    ('Centro:', 'IES Enric Soler i Godes'),
    ('Curso:', '2025-2026'),
    ('URL proyecto:', 'https://sofia-solutions.serveousercontent.com'),
]
for label, val in info:
    p=doc.add_paragraph()
    p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r1=p.add_run(label+' '); r1.bold=True; r1.font.name='Verdana'; r1.font.size=Pt(12)
    r2=p.add_run(val); r2.font.name='Verdana'; r2.font.size=Pt(12)
pb(doc)

print("Portada OK")
doc.save(OUTPUT)
print(f"Guardado parcial en {OUTPUT}")
