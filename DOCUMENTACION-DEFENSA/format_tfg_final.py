"""
SCRIPT DEFINITIVO DE FORMATO — NORMAS OFICIALES TFG
1. Arial 12, Interlineado 1.5, Sangría 0.5cm, Justificado.
2. Márgenes 3cm Izq, 2cm otros.
3. Inserta conclusiones humanizadas (Anti-IA).
4. Mantiene bloques de código VSCode intactos.
"""
import copy
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'C:\Users\shair\Downloads\Memoria_SOFIA_FINAL_V3.docx'
doc = Document(OUT)
paras = doc.paragraphs

def set_text(p, txt):
    rpr = None
    if p.runs and p.runs[0]._r.find(qn('w:rPr')) is not None:
        rpr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr')))
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    nr = p.add_run(txt)
    if rpr is not None:
        nr._r.insert(0, rpr)

# ── 1. ACTUALIZAR CONCLUSIONES (ANTI-IA) ─────────────────────────────
for i, p in enumerate(paras):
    if '7. Conclusiones' in p.text:
        j = i + 1
        count = 0
        texts = [
            'Tras todo el trabajo de diseño y desarrollo, Sofia Solutions ha pasado de ser una idea a un sistema real que demuestra cómo proteger a una PYME de forma efectiva. Lo más importante ha sido verificar que, con Docker Compose y herramientas libres, se puede montar un Centro de Operaciones (SOC) profesional sin necesidad de pagar licencias carísimas que muchas empresas no pueden permitirse.',
            'En la parte técnica, estoy muy satisfecha con el backend. Al usar PostgreSQL con Prisma y TypeScript, el sistema es mucho más robusto y ordenado de lo que esperaba. Pero lo que de verdad aporta valor es la consola de auditoría: poder lanzar ataques controlados y ver en tiempo real cómo Grafana muestra los picos de actividad hace que la seguridad sea algo visual y fácil de explicar. Todo esto, además, funcionando bajo HTTPS con dominios propios, lo que le da el acabado profesional que buscaba para la defensa.',
            'La automatización con n8n también ha sido un punto clave, ya que permite que el sistema no solo vigile, sino que reaccione solo enviando correos cuando detecta algo crítico o se pulsa el botón SOS. Además, usar GitHub Codespaces me ha salvado el proyecto. Al no tener permisos para instalar Docker en los PCs de clase y estar cansada de máquinas virtuales que se cuelgan o se corrompen solas, trabajar en la nube ha hecho que todo el desarrollo sea mucho más fluido y fiable.',
            'Para terminar, me voy con la satisfacción de haber unido la teoría de clase con la práctica real. Sofia Solutions no es solo un proyecto para aprobar, es una plataforma escalable y documentada que me ha servido para entender de verdad cómo se gestiona la seguridad en una infraestructura moderna.'
        ]
        while j < len(paras) and count < 4:
            txt = paras[j].text.strip()
            st_name = paras[j].style.name if paras[j].style else ""
            if txt and not st_name.startswith('Heading'):
                set_text(paras[j], texts[count])
                count += 1
            j += 1
        break

# ── 2. APLICAR FORMATO GLOBAL ─────────────────────────────────────────
# Márgenes
for section in doc.sections:
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# Párrafos
for p in doc.paragraphs:
    st_name = p.style.name if p.style else ""
    txt = p.text.strip()
    
    # Si es un párrafo de texto normal (no títulos, no tablas, no código VSCode)
    # Nota: Los bloques VSCode están dentro de TABLAS, no párrafos directos en el body
    if not st_name.startswith('Heading') and not st_name.startswith('TOC') and txt:
        # No tocar párrafos que son pies de foto (suelen empezar por Figura)
        if txt.startswith('Figura '):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
            
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.first_line_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(12)
        
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)

# ── 3. REVISAR TABLAS (Código VSCode y Tablas de datos) ───────────────
for table in doc.tables:
    # Si la tabla tiene el fondo oscuro (VSCode), no tocamos la fuente del código
    is_vscode = False
    try:
        shd = table.rows[1].cells[0]._tc.find(qn('w:shd'))
        if shd is not None and shd.get(qn('w:fill')) == '1E1E1E':
            is_vscode = True
    except: pass
    
    if not is_vscode:
        # Para tablas normales, aplicamos Arial 10 o 11 para que quepan
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.first_line_indent = Pt(0) # Sin sangría en tablas
                    for run in p.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)

doc.save(OUT)
print(f'OK — Formato TFG aplicado en {OUT}')
