"""Script que inserta bloques VSCode en los Heading3 de 4.5 y mejora las tablas."""
import copy, shutil
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r'C:\Users\shair\Downloads\Memoria_v2.0_Sofia.docx'
OUT = SRC

doc = Document(SRC)

# ── helper: insertar tabla VSCode después de un párrafo ──────────────────
def vscode_block(doc, after_para, filename, code_text):
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Normal Table'
    # Fila título (barra azul VSCode)
    tc0 = tbl.rows[0].cells[0]
    tc0.paragraphs[0].clear()
    shd0 = OxmlElement('w:shd'); shd0.set(qn('w:fill'), '007ACC')
    tc0._tc.get_or_add_tcPr().append(shd0)
    r0 = tc0.paragraphs[0].add_run('  ● ● ●   ' + filename)
    r0.font.name = 'Consolas'; r0.font.size = Pt(9); r0.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    # Fila código (fondo oscuro)
    tc1 = tbl.rows[1].cells[0]
    tc1.paragraphs[0].clear()
    shd1 = OxmlElement('w:shd'); shd1.set(qn('w:fill'), '1E1E1E')
    tc1._tc.get_or_add_tcPr().append(shd1)
    for line in code_text.split('\n'):
        p = tc1.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(13)
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '1E1E1E')
        p._p.get_or_add_pPr().append(shd)
        run = p.add_run(line)
        run.font.name = 'Consolas'; run.font.size = Pt(9)
        # colorear palabras clave básicas
        kws = ['export','async','function','const','let','return','if','await',
               'services:','image:','build:','environment:','volumes:','model','enum']
        color = RGBColor(0xD4,0xD4,0xD4)
        for kw in kws:
            if line.strip().startswith(kw):
                color = RGBColor(0x56,0x9C,0xD6); break
        if line.strip().startswith('#') or line.strip().startswith('//'):
            color = RGBColor(0x6A,0x99,0x55)
        if "'" in line or '"' in line:
            color = RGBColor(0xCE,0x91,0x78)
        run.font.color.rgb = color
    # Borrar el párrafo placeholder vacío que ya hay debajo del heading
    tbl._tbl.getparent().remove(tbl._tbl)
    after_para._p.addnext(tbl._tbl)
    # Añadir línea vacía después
    spacer = OxmlElement('w:p')
    tbl._tbl.addnext(spacer)

# Código para cada sección
CODES = {
    'Docker Compose': ('docker-compose.yml', """services:
  backend:
    build: ./sofia-backend
    environment:
      APP_MODE: ${APP_MODE:-secure}
      DATABASE_URL: postgresql://postgres:postgres@postgres:5432/sofia_solutions
      N8N_WEBHOOK_URL: http://n8n:5678/webhook/alert
    depends_on: [postgres]
  prometheus:
    image: prom/prometheus:latest
    volumes:
      - ./prometheus.yml:/etc/prometheus/prometheus.yml
  tunnel:
    image: alpine
    command: ssh -R sofia-solutions:80:frontend:80 serveo.net"""),
    'Esquema de base de datos': ('prisma/schema.prisma', """model User {
  id           Int      @id @default(autoincrement())
  email        String   @unique
  passwordHash String
  role         Role     @default(CLIENT)
  createdAt    DateTime @default(now())
}
model Ticket {
  id        Int      @id @default(autoincrement())
  titulo    String
  prioridad Priority @default(MEDIA)
  clienteId Int
  cliente   Client   @relation(fields: [clienteId], references: [id])
}
enum Role { ADMIN CLIENT }
enum Priority { BAJA MEDIA ALTA CRITICA }"""),
    'Middleware de autenticaci': ('src/middleware/auth.ts', """export function requireAuth(req: Request, res: Response, next: NextFunction) {
  const isAdmin = req.path.startsWith('/api/admin');
  const suffix  = isAdmin ? '_ADMIN' : '_CLIENT';
  const token   = req.cookies[`accessToken${suffix}`]
                || req.cookies['accessToken']
                || req.headers.authorization?.split(' ')[1];
  if (!token) throw new ApiError(401, 'Authentication required');
  const payload = verifyAccessToken(token);
  req.user = payload;
  next();
}"""),
    'Sistema de doble modo': ('src/controllers/auth.controller.ts', """// MODO VULNERABLE — permite SQL Injection bypass
export async function loginV1(req: Request, res: Response) {
  const { email, password } = req.body;
  const users: any[] = await prisma.$queryRawUnsafe(
    `SELECT * FROM "User" WHERE "email" = '${email}'`
  );
  const isInjection = email.includes("'") || email.includes("--");
  if (!users[0] || (!isInjection && users[0].password !== password))
    throw new ApiError(401, 'Invalid credentials (Vulnerable Mode)');
  res.json({ accessToken: signAccessToken({ sub: users[0].id }) });
}

// MODO SEGURO — Bcrypt + Prisma parametrizado (inmune a SQLi)
export async function loginV2(req: Request, res: Response) {
  const { email, password } = req.body;
  const user = await prisma.user.findUnique({ where: { email } });
  if (!user || !(await bcrypt.compare(password, user.passwordHash)))
    throw new ApiError(401, 'Invalid identity credentials');
  res.json({ accessToken: signAccessToken({ sub: user.id, role: user.role }) });
}"""),
    'Script de despliegue': ('scripts/SOFIA-INICIAR.sh', """#!/bin/bash
set -e
echo "🔒 Sofia Solutions — Iniciando stack completo..."
APP_MODE=${APP_MODE:-secure}
docker compose down --remove-orphans
docker compose up -d --build
echo ""
echo "✅ Stack activo. Servicios disponibles:"
echo "   🌐 Frontend  → http://localhost:8000"
echo "   📊 Grafana   → http://localhost:3000"
echo "   🔗 n8n       → http://localhost:5678"
docker compose ps"""),
    'Panel SOC': ('frontend-php/public/assets/soc.js', """// SOC Monitor — Actualización automática cada 8 segundos
async function refreshSOCMetrics() {
  const token = sessionStorage.getItem('accessToken_ADMIN');
  const res = await fetch('/api/admin/metrics/summary', {
    headers: { 'Authorization': `Bearer ${token}` }
  });
  const data = await res.json();
  document.getElementById('total-events').textContent    = data.totalEvents;
  document.getElementById('blocked-attacks').textContent = data.blockedAttacks;
  document.getElementById('active-sessions').textContent = data.activeSessions;
  updateChart(data.timeline);
}

setInterval(refreshSOCMetrics, 8000);
refreshSOCMetrics();"""),
}

# Insertar bloques después de cada Heading3 correspondiente
for i, p in enumerate(doc.paragraphs):
    try:
        sn = p.style.name or ''
    except:
        sn = ''
    if sn == 'Heading 3':
        txt = p.text.strip()
        for key, (fname, code) in CODES.items():
            if key.lower() in txt.lower():
                vscode_block(doc, p, fname, code)
                break

# ── Mejorar estilos de tablas existentes ────────────────────────────────
from docx.oxml.ns import qn as ns_qn
HEADER_COLOR = '0E4D92'
ALT_COLOR    = 'EBF2FA'

for table in doc.tables:
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            shd = OxmlElement('w:shd')
            if r_idx == 0:
                shd.set(qn('w:fill'), HEADER_COLOR)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                        run.bold = True
            elif r_idx % 2 == 0:
                shd.set(qn('w:fill'), ALT_COLOR)
            else:
                shd.set(qn('w:fill'), 'FFFFFF')
            cell._tc.get_or_add_tcPr().append(shd)

doc.save(OUT)
print(f'OK — Guardado: {OUT}')
