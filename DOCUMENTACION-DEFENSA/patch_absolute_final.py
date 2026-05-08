"""
ULTIMATE PATCH — USES Sofia_Solutions_Memoria_TFG.docx AS BASE.
"""
import copy, shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r'C:\Users\shair\Downloads\Sofia_Solutions_Memoria_TFG.docx'
OUT = r'C:\Users\shair\Downloads\Memoria_SOFIA_FINAL_V3.docx'

shutil.copy2(SRC, OUT)
doc = Document(OUT)
paras = doc.paragraphs

def set_text(p, txt):
    rpr = None
    if p.runs and p.runs[0]._r.find(qn('w:rPr')) is not None:
        rpr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr')))
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    nr = p.add_run(txt)
    if rpr is not None:
        nr._r.insert(0, rpr)

def insert_placeholder_after(para_element, label):
    new_p = OxmlElement('w:p')
    orig_pPr = para_element.find(qn('w:pPr'))
    if orig_pPr is not None: new_p.append(copy.deepcopy(orig_pPr))
    r = OxmlElement('w:r')
    t = OxmlElement('w:t'); t.text = f'(Insertar aquí captura: {label})'
    r.append(t)
    new_p.append(r)
    para_element.addnext(new_p)
    return new_p

# 1. Buscar Conclusiones
found_conc = False
for i, p in enumerate(paras):
    if '7. Conclusiones' in p.text:
        found_conc = True
        j = i + 1
        count = 0
        while j < len(paras) and count < 4:
            txt = paras[j].text.strip()
            # Verificamos si el estilo existe antes de acceder a .name
            style_name = paras[j].style.name if paras[j].style else ""
            if txt and not style_name.startswith('Heading'):
                if count == 0:
                    set_text(paras[j], 'Poder ver en tiempo real lo que pasa en el sistema gracias a las gráficas de Grafana y Prometheus ha sido clave. Ya no son solo datos aburridos que nadie lee, sino que ves el ataque ocurriendo en directo y puedes reaccionar. Además, n8n hace que el sistema responda por sí solo: si alguien intenta entrar a la fuerza o si el cliente pulsa el botón de emergencia SOS, llega un correo de alerta al SOC de forma inmediata.')
                elif count == 1:
                    set_text(paras[j], 'El uso de GitHub Codespaces ha sido un gran acierto, porque me ha permitido olvidarme de configurar el PC. Simplemente abres el navegador y ya tienes todo listo para funcionar. Gracias a que usamos Docker, el proyecto se puede desplegar en segundos en cualquier ordenador y sé que funcionará exactamente igual que en mi equipo, lo que da mucha tranquilidad para el día de la defensa.')
                elif count == 2:
                    set_text(paras[j], 'En cuanto a las herramientas, GitHub Codespaces ha sido una de las mejores decisiones. Poder tener todo el entorno configurado en la nube me ha permitido trabajar desde cualquier sitio sin preocuparme de dependencias. Grafana y Prometheus me han sorprendido mucho: al principio pensaba que serían muy complejas, pero ver cómo pueden transformar los logs en gráficas que se mueven en directo hace que la seguridad sea mucho más fácil de entender y de explicar.')
                elif count == 3:
                    set_text(paras[j], 'Finalmente, Sofia Solutions se posiciona como un ecosistema de defensa profesional que cierra la brecha entre la teoría académica y la operación de seguridad real. El resultado es una plataforma escalable, documentada y preparada para enfrentar los ataques más comunes de la actualidad.')
                count += 1
            j += 1
        break

# 2. Fase 5 n8n
for i, p in enumerate(paras):
    if 'Fase 5' in p.text:
        for j in range(i+1, min(i+20, len(paras))):
            if 'Grafana fue lo más visual' in paras[j].text:
                new_p = OxmlElement('w:p')
                orig_pPr = paras[j]._p.find(qn('w:pPr'))
                if orig_pPr is not None: new_p.append(copy.deepcopy(orig_pPr))
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = ('n8n ha sido la pieza que ha permitido automatizar las respuestas. He configurado un flujo que recibe los avisos del backend y los convierte en correos electrónicos reales. De esta forma, si ocurre algo grave como un ataque de fuerza bruta o si el usuario pulsa el botón de emergencia (SOS), el equipo de seguridad recibe un aviso inmediato.')
                r.append(t)
                new_p.append(r)
                paras[j]._p.addnext(new_p)
                insert_placeholder_after(new_p, 'Flujo n8n: Alertas SOC')
                break
        break

doc.save(OUT)
print(f'OK — Documento FINAL guardado en: {OUT}')
