"""
Patch definitivo — actualiza texto en Sofia_Solutions_Memoria_TFG.docx
manteniendo TODO el formato, imágenes y estilos intactos.
Guarda como Memoria_v2.0_Sofia.docx en Descargas.
"""
import copy, shutil
from docx import Document
from docx.oxml.ns import qn

SRC = r'C:\Users\shair\Downloads\Sofia_Solutions_Memoria_TFG.docx'
OUT = r'C:\Users\shair\Downloads\Memoria_v2.0_Sofia.docx'

shutil.copy2(SRC, OUT)
doc = Document(OUT)
paras = doc.paragraphs

def set_text(p, txt):
    """Reemplaza solo el texto del párrafo conservando el estilo de fuente."""
    if not p.runs:
        p.add_run(txt)
        return
    # Guardar rPr del primer run
    rpr = None
    if p.runs[0]._r.find(qn('w:rPr')) is not None:
        rpr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr')))
    # Borrar todos los runs
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    # Nuevo run
    nr = p.add_run(txt)
    if rpr is not None:
        nr._r.insert(0, rpr)

def set_range(start, end, texts):
    """Rellena párrafos no vacíos en rango con los textos dados."""
    targets = []
    for i in range(start, min(end, len(paras))):
        p = paras[i]
        try:
            sname = p.style.name or ''
        except:
            sname = ''
        if sname.startswith('Heading'):
            continue
        txt = p.text.strip()
        if not txt:
            continue
        if p._p.findall('.//' + qn('w:drawing')):
            continue
        targets.append(p)
    for i, new_txt in enumerate(texts):
        if i < len(targets):
            set_text(targets[i], new_txt)

def heading_range(keyword):
    """Devuelve (start+1, next_heading_idx) para un heading que contenga keyword."""
    kw = keyword.lower()
    for i, p in enumerate(paras):
        try:
            sname = p.style.name or ''
        except:
            sname = ''
        if sname.startswith('Heading') and kw in p.text.lower():
            for j in range(i+1, len(paras)):
                try:
                    sn2 = paras[j].style.name or ''
                except:
                    sn2 = ''
                if sn2.startswith('Heading') and paras[j].text.strip() and j != i+1:
                    return i+1, j
            return i+1, len(paras)
    return None, None

# ════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN
# ════════════════════════════════════════════════════════════
s, e = heading_range('introducci')
if s:
    set_range(s, e, [
        'Al fijarnos en el panorama empresarial de hoy, se ve claro que muchísimas PYMEs tienen un problema de seguridad enorme porque, sencillamente, no pueden permitirse grandes inversiones en infraestructura de ciberseguridad. Contratar un SOC tradicional puede costar decenas de miles de euros al año, algo inasumible para una empresa pequeña. Por eso quise centrar mi proyecto en algo útil de verdad: montar un entorno completo que simule lo que necesitaría una PYME real para gestionar su seguridad IT, usando exclusivamente software libre y contenedores Docker.',
        'La plataforma Sofia Solutions es el resultado de ese trabajo. Es un sistema web corporativo con dos partes: un frontend en PHP con panel de administración para gestionar clientes, servicios y tickets; y un backend en Node.js con TypeScript que incorpora una API capaz de funcionar en dos modos —vulnerable y seguro— para hacer demostraciones educativas de ataques reales y sus contramedidas. Todo se complementa con Grafana y Prometheus para monitorizar el sistema en tiempo real, y con n8n para la automatización de alertas de seguridad.',
        'Módulos del ciclo aplicados: Seguridad y Alta Disponibilidad (base para entender vectores OWASP y defensas), Implantación de Aplicaciones Web (frontend PHP/HTML/CSS con diseño Cyberpunk-Glassmorphism), Administración de Sistemas Operativos (scripts Bash y PowerShell de automatización), Servicios de Red e Internet (Docker Compose, exposición de puertos, túnel SSH a serveo.net), y Administración de SGBD (diseño E-R PostgreSQL con Prisma ORM y migraciones).',
        'El proyecto ha evolucionado bastante desde la idea inicial. Al principio solo quería montar un entorno vulnerable para practicar ataques OWASP, pero se fue convirtiendo en algo más ambicioso: una plataforma completa con panel SOC, gestión de clientes, sistema de tickets, monitorización real con Grafana y automatización de alertas con n8n.',
    ])

# ════════════════════════════════════════════════════════════
# 2. OBJETIVOS
# ════════════════════════════════════════════════════════════
s, e = heading_range('objectius')
if not s: s, e = heading_range('2. obj')
if s:
    set_range(s, e, [
        'Objetivo principal: montar un laboratorio de ciberseguridad funcional que sirva como plataforma de demostración para una PYME, comparando una infraestructura sin protecciones frente a una correctamente asegurada.',
        'Diseñar e implementar una arquitectura de microservicios con Docker Compose donde frontend, backend, base de datos, monitorización y automatización estén aislados en contenedores independientes.',
        'Crear una API REST en Node.js con TypeScript con sistema de doble modo (APP_MODE=secure/vulnerable) que permita hacer demos educativas de las vulnerabilidades más críticas del OWASP Top 10.',
        'Desarrollar un frontend corporativo en PHP con panel de administración completo para gestionar clientes, servicios IT y tickets de incidencias de seguridad.',
        'Implementar un panel SOC en tiempo real con Grafana y Prometheus más un feed de actividad actualizado cada 8 segundos mediante polling JavaScript.',
        'Configurar acceso HTTPS público mediante túnel SSH a serveo.net sin necesidad de IP fija ni VPS contratado.',
        'Automatizar el despliegue completo con scripts Bash y PowerShell de forma que un solo comando levante todo el stack.',
        'Aplicar medidas de seguridad reales en modo seguro: Bcrypt con 12 rounds, JWT, cabeceras CSP, rate limiting de 5 intentos por IP y consultas parametrizadas con Prisma ORM.',
    ])

# ════════════════════════════════════════════════════════════
# 3. ANÁLISIS REQUERIMIENTOS — 3.1
# ════════════════════════════════════════════════════════════
s, e = heading_range('3.1')
if not s: s, e = heading_range('situaci')
if s:
    set_range(s, e, [
        'El cliente es Sofia Solutions, una PYME ficticia de servicios IT que representa el caso de uso real del proyecto. La situación inicial era la habitual en muchas empresas pequeñas: gestión con hojas de cálculo y correos electrónicos, sin sistema centralizado, sin trazabilidad de incidentes y sin ningún tipo de monitorización de seguridad. Partíamos de cero, sin ningún servidor configurado.',
        'Los requerimientos funcionales identificados fueron: panel web para gestión de clientes y servicios IT, sistema de tickets para incidencias con priorización BAJA/MEDIA/ALTA/CRITICA, API de autenticación con roles ADMIN y CLIENT, panel SOC con feed de actividad en tiempo real, módulo de demo educativa de vulnerabilidades en modo dual, y tres planes de servicio: Individual 499€/mes, Business 1.500€/mes y Business Max 4.200€/mes.',
        'Los requerimientos no funcionales incluyeron: despliegue en menos de 5 minutos desde cero con Docker Compose, acceso HTTPS mediante túnel serveo.net, portabilidad total a cualquier Linux con Docker instalado, y monitorización continua con alertas automáticas vía n8n.',
    ])

# ════════════════════════════════════════════════════════════
# 3.2 ALTERNATIVAS
# ════════════════════════════════════════════════════════════
s, e = heading_range('3.2')
if not s: s, e = heading_range('alternatives')
if s:
    set_range(s, e, [
        'Para el backend se valoraron WordPress con plugins (descartado por falta de control fino sobre la API y dificultad para dockerizar), Django en Python (descartado por curva de aprendizaje alta) y Laravel (descartado por ser más pesado y no permitir fácilmente el sistema de doble modo). Se eligió Node.js con Express y TypeScript por ser ligero, tener tipado estático fuerte y permitir el patrón APP_MODE de forma limpia.',
        'Para la base de datos se valoraron MySQL (licencia dual Oracle), SQLite (no apta para múltiples conexiones simultáneas) y MongoDB (no relacional, no apropiada para datos cliente/ticket relacionales). Se eligió PostgreSQL 15 por ser 100% open source, robusta en entornos multi-conexión y tener integración excelente con Prisma ORM.',
        'Para la monitorización se valoraron Zabbix (muy pesado, difícil de dockerizar), ELK Stack (requiere demasiada RAM para el alcance del proyecto) y Datadog (SaaS de pago). Se eligió Prometheus más Grafana por ser ligeros, dockerizables y generar dashboards modernos con queries PromQL.',
    ])

# ════════════════════════════════════════════════════════════
# 3.3 PRESUPUESTO
# ════════════════════════════════════════════════════════════
s, e = heading_range('3.3')
if not s: s, e = heading_range('pressupost')
if not s: s, e = heading_range('presupuesto')
if s:
    set_range(s, e, [
        'La infraestructura elegida tiene un coste mínimo: VPS Hetzner CX21 (2 vCPU, 4 GB RAM) a 5,83€/mes, dominio .com a 0,92€/mes y certificado SSL con Let\'s Encrypt gratuito. Total de infraestructura: 7€/mes, 84€/año.',
        'El coste de desarrollo valorado a precio de mercado (25€/hora) sería: análisis y diseño 20h (500€), backend API 40h (1.000€), frontend PHP 30h (750€), Docker y automatización 15h (375€), monitorización 10h (250€), pruebas y documentación 20h (500€). Total: 135 horas, 3.375€.',
        'Comparativa: una solución propietaria equivalente con Splunk SIEM más licencias de seguridad tendría un coste de entre 15.000 y 40.000€ anuales solo en licencias, sin contar el hardware ni el personal especializado.',
    ])

# ════════════════════════════════════════════════════════════
# 4.5 CODIFICACIÓN
# ════════════════════════════════════════════════════════════
s, e = heading_range('codific')
if s:
    set_range(s, e, [
        'La programación está organizada por capas: el frontend PHP consume la API REST del backend mediante fetch con el token JWT en la cabecera Authorization. Lo más representativo del código son el controlador de autenticación dual (loginV1 vulnerable vs loginV2 seguro), el middleware de detección de ataques con bypass para la consola de auditoría, el rate limiter con alerta automática a n8n, el esquema Prisma con los modelos de datos, y los scripts de automatización Bash y PowerShell.',
    ])

# ════════════════════════════════════════════════════════════
# 4.6 ADMINISTRACIÓN
# ════════════════════════════════════════════════════════════
s, e = heading_range('administrac')
if s:
    set_range(s, e, [
        'La gestión del sistema opera a dos niveles. A nivel de infraestructura, Docker Compose permite arrancar, parar o reiniciar cada servicio de forma independiente sin afectar al resto: docker compose up -d para levantar todo, docker compose logs -f backend para ver logs en tiempo real, y APP_MODE=vulnerable docker compose up -d para cambiar al modo de demo. Prisma Studio en http://localhost:5556 proporciona una interfaz visual de la base de datos muy útil para las demos. A nivel de aplicación, el panel admin gestiona clientes, tickets, servicios y el SOC. Roles: ADMIN (acceso total) y CLIENT (solo su dashboard y tickets propios). Credenciales de demo: admin@sofia.local / S0f1a_Secur3!_2026.',
    ])

# ════════════════════════════════════════════════════════════
# 4.7 SEGURIDAD
# ════════════════════════════════════════════════════════════
s, e = heading_range('seguretat')
if not s: s, e = heading_range('seguridad')
if s:
    set_range(s, e, [
        'La seguridad está implementada tanto en el código como en la infraestructura de red, con un modelo de doble modo que permite comparar el sistema sin protecciones con el sistema correctamente asegurado. En modo seguro: SQLi se bloquea con consultas parametrizadas de Prisma ORM (findUnique tipado); la fuerza bruta se limita a 5 intentos por IP cada 15 minutos con alerta automática a n8n; las contraseñas se hashean con Bcrypt a 12 rounds ($2b$12$...); XSS se previene con htmlspecialchars() en PHP y cabecera Content-Security-Policy; IDOR se bloquea verificando que el usuario autenticado sea el propietario del recurso. A nivel de infraestructura, PostgreSQL no expone su puerto al exterior: solo es accesible desde la red interna Docker sofia-network. El túnel serveo.net proporciona certificado SSL/TLS automático para todo el tráfico exterior, garantizando HTTPS end-to-end.',
    ])

# ════════════════════════════════════════════════════════════
# FASES
# ════════════════════════════════════════════════════════════
s, e = heading_range('fase 1')
if not s: s, e = heading_range('planificaci')
if s:
    set_range(s, e, [
        'Empecé definiendo los objetivos basándome en el temario de ASIX e investigué las vulnerabilidades más críticas del OWASP Top 10 para representarlas de forma didáctica. La decisión más importante de esta fase fue el sistema de doble modo: en lugar de dos proyectos separados, una variable de entorno APP_MODE hace que el mismo código se comporte de forma diferente, sin duplicar nada. Esta decisión ahorró semanas de trabajo.',
    ])

s, e = heading_range('fase 2')
if not s: s, e = heading_range('disseny')
if s:
    set_range(s, e, [
        'Dibujé el esquema E-R en papel antes de pasarlo a Prisma y tuve que rehacerlo dos veces. El primer diseño conectaba los tickets directamente a los servicios, pero era más correcto conectarlos a los clientes con técnicos (usuarios admin) asignados. Ese cambio lo hice tarde y me costó actualizar bastantes queries. También el healthcheck de Docker fue problemático al principio: Prisma intentaba conectarse a PostgreSQL antes de que terminara de arrancar, así que añadí sleep 15 en el script de inicio como solución provisional.',
    ])

s, e = heading_range('fase 3')
if not s: s, e = heading_range('backend (api)')
if s:
    set_range(s, e, [
        'Desarrollé el servidor con Node.js y TypeScript programando el sistema de doble modo con defensas activas o desactivadas según APP_MODE. La fase más larga porque nunca había trabajado con TypeScript de forma seria. El compilador interrumpía el build del Dockerfile por noImplicitAny; ajusté tsconfig.json para priorizar que el stack arrancara sobre el linting defensivo. También tuve el error de los retornos de carro DOS (\\r\\n) en los scripts: el intérprete Alpine los leía como "illegal option -\\r". Solución: convertir a saltos de línea Unix antes del build con sed.',
    ])

s, e = heading_range('fase 4')
if not s: s, e = heading_range('frontend corporat')
if s:
    set_range(s, e, [
        'Usé lo aprendido en Implantación de Aplicaciones Web para desarrollar la interfaz en PHP con estilo Cyberpunk/Glassmorphism: fondo oscuro, efectos de cristal con backdrop-filter blur, colores neón cyan y violeta, animaciones CSS. El problema más curioso fue la integración PHP-Express: el navegador enviaba cookies de sesión PHP pero el backend esperaba JWT en la cabecera Authorization. Tuve que hacer que PHP extrajera el token JWT de $_SESSION y lo incluyera en cada fetch al backend.',
    ])

s, e = heading_range('fase 5')
if not s: s, e = heading_range('prometheus')
if s:
    set_range(s, e, [
        'Configurar Prometheus fue bastante más manual de lo que esperaba. Tuve que añadir endpoints /metrics en la API con prom-client para exponer la métrica sofia_login_attempts_total etiquetada por mode y result. Prometheus hace scraping cada 15 segundos.',
        'Grafana fue lo más visual del proyecto. Una vez conectado a Prometheus con la query PromQL correcta, las gráficas de intentos de login muestran en tiempo real el pico exacto de cuando se ejecuta el ataque de fuerza bruta. También configuré n8n con un workflow Webhook→Set→Send Email para las alertas SOS, aunque el SMTP tardó en funcionar porque n8n necesita la variable WEBHOOK_URL apuntando al dominio público, no a localhost.',
    ])

s, e = heading_range('fase 6')
if not s: s, e = heading_range('pruebas')
if s:
    set_range(s, e, [
        'Las pruebas las hice principalmente con los scripts propios del repositorio (SOFIA-ATAQUES.sh y sofia-audit-framework.py) y con Postman para los endpoints de la API. La parte más interesante fue verificar el sistema de doble modo: cambiar APP_MODE=vulnerable, ejecutar el SQLi bypass con \' OR \'1\'=\'1\'-- y ver el token JWT devuelto; luego cambiar a APP_MODE=secure y verificar que el mismo payload devuelve 403. También verifiqué que el rate limiting bloqueaba en HTTP 429 tras 5 intentos y que Grafana mostraba el pico de intentos en tiempo real durante el brute force.',
    ])

# ════════════════════════════════════════════════════════════
# 6. AMPLIACIONES
# ════════════════════════════════════════════════════════════
s, e = heading_range('6.1')
if s:
    set_range(s, e, [
        'Actualmente el sistema en modo vulnerable permite demostrar inyección SQL (auth bypass y UNION exfiltration), fuerza bruta de diccionario, IDOR en tickets de soporte, XSS reflejado y DoS por inundación HTTP. La ampliación sería completar la matriz OWASP Top 10 añadiendo SSRF (Server-Side Request Forgery), vulnerabilidades de deserialización insegura y Path Traversal avanzado con lectura de archivos del sistema, cada una con su contramedida activa en modo seguro para que la demo comparativa sea completa.',
        'El coste de esta ampliación se estima en unas 40 horas adicionales de desarrollo a precio de mercado (1.000€), incluyendo documentación de cada vector de ataque.',
    ])

s, e = heading_range('6.2')
if s:
    set_range(s, e, [
        'El panel SOC actual muestra métricas básicas con Chart.js y polling cada 8 segundos. Una ampliación importante sería integrarlo con Loki para centralizar los logs de todos los contenedores Docker y permitir búsquedas históricas. Otra mejora sería añadir un mapa de IPs en tiempo real con MaxMind GeoIP que muestre geográficamente de dónde provienen los ataques, algo que daría mucho impacto visual en demos ante clientes.',
        'Coste estimado: 24 horas adicionales (600€) para la integración de Loki y el componente de mapa geográfico.',
    ])

s, e = heading_range('6.3')
if s:
    set_range(s, e, [
        'Ya existe scripts/sofia-audit-framework.py (Python interactivo en terminal). La ampliación más valiosa sería integrarlo directamente en la consola web /consola de forma que los módulos de auditoría (SQLi, Brute Force, IDOR, DoS, XSS) sean ejecutables desde el navegador con log detallado de cada payload enviado y respuesta recibida, y que el sistema pueda exportar el informe de auditoría completo en PDF.',
        'Esta ampliación permitiría ofrecer el módulo de auditoría como servicio a clientes de la plataforma, aumentando el valor del producto. Coste estimado: 20 horas (500€).',
    ])

s, e = heading_range('6.4')
if s:
    set_range(s, e, [
        'Una ampliación a más largo plazo sería desarrollar una app móvil sencilla con React Native para que los técnicos puedan consultar y actualizar tickets asignados desde el teléfono, recibir notificaciones push de alertas SOS y ver el estado del SOC. La API REST ya está documentada y lista para consumirse desde mobile. Estimación: 80 horas adicionales (2.000€).',
    ])

# ════════════════════════════════════════════════════════════
# 7. CONCLUSIONES
# ════════════════════════════════════════════════════════════
s, e = heading_range('7. conclusi')
if not s: s, e = heading_range('conclusi')
if s:
    set_range(s, e, [
        'Tras el proceso de diseño, desarrollo e implementación de la plataforma Sofia Solutions, se ha logrado consolidar un sistema que no solo cumple con los objetivos técnicos iniciales, sino que también aporta una visión práctica sobre la gestión de la ciberseguridad en la pequeña y mediana empresa. Se ha verificado la viabilidad de democratizar los servicios de un Centro de Operaciones de Seguridad (SOC) mediante el uso de tecnologías de código abierto e infraestructuras basadas íntegramente en Docker. El uso de Docker Compose ha permitido orquestar servicios complejos de forma eficiente, garantizando una monitorización crítica sin incurrir en los elevados costes de las licencias propietarias tradicionales.',
        'Un punto fundamental ha sido la profesionalización del backend y la gestión de la persistencia con PostgreSQL y el ORM Prisma, lo que ha permitido establecer un entorno multi-cliente robusto y optimizado. Asimismo, el desarrollo de la consola de auditoría técnica ha permitido recrear de forma inmersiva el flujo de trabajo de un experto, facilitando la ejecución controlada de exploits de inyección SQL, ataques de fuerza bruta y simulaciones de compromiso de sistemas. Esta faceta ofensiva ha sido clave para validar la robustez de la plataforma, que queda expuesta públicamente con dominio personalizado y certificado SSL/TLS, garantizando que el tráfico viaje cifrado mediante HTTPS end-to-end.',
        'La integración de la telemetría mediante Grafana y Prometheus ha permitido transformar datos brutos en inteligencia accionable, visualizando el ciclo de vida de los incidentes a través de dashboards interactivos. Esta capacidad de análisis se ve reforzada por la automatización mediante n8n para la gestión de alertas SOS, garantizando una respuesta inmediata ante eventos críticos. Además, la adopción de GitHub Codespaces ha eliminado la necesidad de máquinas virtuales locales pesadas: gracias a la portabilidad intrínseca de las imágenes de Docker, el proyecto puede desplegarse en segundos en un entorno cloud-native, con un comportamiento idéntico independientemente del hardware.',
        'Finalmente, Sofia Solutions se posiciona como un ecosistema de defensa profesional que cierra la brecha entre la teoría académica y la operación de seguridad real. El uso de Docker como estándar de despliegue, junto con la automatización de flujos mediante n8n, la observabilidad con Prometheus y Grafana, y la documentación exhaustiva del código con TypeScript, demuestra que la modernización de infraestructuras es el pilar indispensable de la ciberseguridad corporativa moderna. El resultado es una plataforma escalable, documentada y preparada para enfrentar los vectores de ataque más comunes de la actualidad.',
    ])

# ════════════════════════════════════════════════════════════
# 9. SOPORTE FÍSICO
# ════════════════════════════════════════════════════════════
s, e = heading_range('soporte f')
if not s: s, e = heading_range('suport f')
if s:
    set_range(s, e, [
        'En la contraportada se adjunta un pendrive USB con el contenido completo del proyecto. El pendrive contiene todos los archivos necesarios para desplegar y demostrar Sofia Solutions desde cero en cualquier máquina con Docker instalado.',
    ])

doc.save(OUT)
print(f'OK - Guardado en:\n{OUT}')
