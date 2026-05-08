"""
Actualiza el texto de la memoria manteniendo TODOS los estilos, imágenes y formato original.
Solo reemplaza el contenido de texto en cada sección.
"""
import copy
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

SRC = r'C:\Users\shair\Downloads\Memoria. Entrega 4.docx'
OUT = r'C:\Users\shair\Downloads\Memoria_v2.0_Sofia.docx'

doc = Document(SRC)

def set_para_text(para, new_text):
    """Reemplaza el texto de un párrafo conservando el estilo del primer run."""
    # Guardar formato del primer run si existe
    first_run_xml = None
    if para.runs:
        first_run_xml = copy.deepcopy(para.runs[0]._r)
    # Borrar todos los runs del párrafo
    for r in para.runs:
        r._r.getparent().remove(r._r)
    # Añadir nuevo run con el texto
    new_run = para.add_run(new_text)
    if first_run_xml is not None:
        # Copiar formato de fuente del run original
        orig_rpr = first_run_xml.find(qn('w:rPr'))
        if orig_rpr is not None:
            new_run._r.insert(0, copy.deepcopy(orig_rpr))

def find_section_range(paragraphs, heading_keyword):
    """Devuelve (start, end) índices de párrafos de contenido bajo un heading."""
    start = None
    for i, p in enumerate(paragraphs):
        if p.style.name.startswith('Heading') and heading_keyword.lower() in p.text.lower():
            start = i + 1
        elif start is not None and p.style.name.startswith('Heading') and p.text.strip():
            return start, i
    if start is not None:
        return start, len(paragraphs)
    return None, None

def get_content_paras(paragraphs, start, end):
    """Devuelve pares (índice, párrafo) de párrafos con texto editable en el rango."""
    result = []
    for i in range(start, end):
        p = paragraphs[i]
        # Saltar vacíos, figuras y headings
        text = p.text.strip()
        if not text:
            continue
        if p.style.name.startswith('Heading'):
            continue
        if text.startswith('Figura') or text.startswith('Descripci'):
            continue
        # Saltar si el párrafo tiene imágenes
        if p._p.findall('.//' + qn('w:drawing')):
            continue
        result.append((i, p))
    return result

paras = doc.paragraphs

# ── SECCIÓN 1: INTRODUCCIÓN ────────────────────────────────────────────────
intro_texts = [
    'Al fijarnos en el panorama empresarial de hoy, se ve claro que muchísimas PYMEs tienen un problema de seguridad enorme porque, sencillamente, no pueden permitirse grandes inversiones en infraestructura de ciberseguridad. Contratar un SOC tradicional puede costar decenas de miles de euros al año, algo inasumible para una empresa pequeña. Por eso quise centrar mi proyecto en algo útil de verdad: montar un entorno completo que simule lo que necesitaría una PYME real para gestionar su seguridad IT, usando exclusivamente software libre y contenedores Docker.',
    'La plataforma Sofia Solutions es el resultado de ese trabajo. Es un sistema web corporativo con dos partes: un frontend en PHP con panel de administración para gestionar clientes, servicios y tickets; y un backend en Node.js con TypeScript que incorpora una API capaz de funcionar en dos modos —vulnerable y seguro— para hacer demostraciones educativas de ataques reales y sus contramedidas. Todo se complementa con Grafana y Prometheus para monitorizar el sistema en tiempo real.',
    'Módulos del ciclo aplicados: Seguridad y Alta Disponibilidad (vectores OWASP), Implantación de Aplicaciones Web (frontend PHP/HTML/CSS), Administración de Sistemas Operativos (scripts Bash/PowerShell), Servicios de Red e Internet (Docker, túnel SSL serveo.net) y Administración de SGBD (diseño E-R PostgreSQL, Prisma ORM).',
    'El proyecto ha evolucionado bastante desde la idea inicial. Al principio solo quería montar un entorno vulnerable para practicar, pero se fue convirtiendo en algo más ambicioso: plataforma completa con panel SOC, gestión de clientes, sistema de tickets y monitorización real.',
]
s, e = find_section_range(paras, 'introducci')
if s:
    content = get_content_paras(paras, s, e)
    for idx, (pi, p) in enumerate(content):
        if idx < len(intro_texts):
            set_para_text(p, intro_texts[idx])

# ── SECCIÓN 2: OBJETIVOS ───────────────────────────────────────────────────
obj_texts = [
    'Objetivo principal: montar un laboratorio de ciberseguridad funcional que sirva como plataforma de demostración para una PYME, comparando una infraestructura sin protecciones frente a una correctamente asegurada.',
    '1. Diseñar e implementar arquitectura de microservicios con Docker Compose (frontend, backend, BBDD, monitorización aislados).',
    '2. Crear API REST en Node.js+TypeScript con sistema de doble modo (APP_MODE=secure/vulnerable) para demos educativas de OWASP Top 10.',
    '3. Desarrollar frontend corporativo en PHP con panel admin para gestionar clientes, servicios IT y tickets de incidencias.',
    '4. Implementar panel SOC en tiempo real con Grafana+Prometheus y feed de actividad actualizado cada 8 segundos.',
    '5. Configurar acceso HTTPS público mediante túnel serveo.net sin necesidad de IP fija ni VPS contratado.',
    '6. Automatizar despliegue completo con scripts Bash/PowerShell (un solo comando levanta todo el stack).',
    '7. Aplicar medidas de seguridad reales en modo seguro: Bcrypt (12 rounds), JWT, CSP, rate limiting y consultas parametrizadas con Prisma.',
]
s, e = find_section_range(paras, 'objectius')
if not s:
    s, e = find_section_range(paras, 'objetivos')
if s:
    content = get_content_paras(paras, s, e)
    for idx, (pi, p) in enumerate(content):
        if idx < len(obj_texts):
            set_para_text(p, obj_texts[idx])

# ── SECCIÓN 4.5: CODIFICACIÓN ──────────────────────────────────────────────
cod_texts = [
    'La programación está organizada por capas: el frontend PHP consume la API REST del backend mediante fetch con JWT en la cabecera Authorization. Los fragmentos más representativos son el controlador de autenticación dual, el middleware de detección de ataques y los scripts de automatización de Docker.',
]
s, e = find_section_range(paras, 'codific')
if s:
    content = get_content_paras(paras, s, e)
    for idx, (pi, p) in enumerate(content):
        if idx < len(cod_texts):
            set_para_text(p, cod_texts[idx])

# ── SECCIÓN 4.6: ADMINISTRACIÓN ───────────────────────────────────────────
adm_texts = [
    'La gestión del sistema opera a dos niveles. A nivel de infraestructura, Docker Compose permite arrancar, parar o reiniciar cada servicio de forma independiente sin afectar al resto. A nivel de aplicación, el panel de administración web permite gestionar clientes, tickets, servicios y ver el SOC en tiempo real. Prisma Studio (http://localhost:5556) proporciona una interfaz visual de la base de datos útil para demos. Los roles son ADMIN (acceso total) y CLIENT (solo su dashboard y tickets propios). Credenciales de demo: admin@sofia.local / S0f1a_Secur3!_2026.',
]
s, e = find_section_range(paras, 'administrac')
if s:
    content = get_content_paras(paras, s, e)
    for idx, (pi, p) in enumerate(content):
        if idx < len(adm_texts):
            set_para_text(p, adm_texts[idx])

# ── SECCIÓN 4.7: SEGURIDAD ─────────────────────────────────────────────────
seg_texts = [
    'La seguridad está implementada tanto en el código como en la infraestructura de red. En modo seguro: SQLi → Prisma ORM con consultas parametrizadas; Fuerza Bruta → rate limiting 5 intentos/15 min con alerta automática a n8n; Contraseñas → Bcrypt 12 rounds ($2b$12$...); XSS → htmlspecialchars() PHP + cabecera CSP; IDOR → middleware que verifica que el usuario autenticado sea el propietario del recurso. PostgreSQL no expone puerto al exterior: solo es accesible desde la red interna Docker. El túnel serveo.net proporciona SSL/TLS automático para todo el tráfico exterior.',
]
s, e = find_section_range(paras, 'seguretat')
if not s:
    s, e = find_section_range(paras, 'seguridad')
if s:
    content = get_content_paras(paras, s, e)
    for idx, (pi, p) in enumerate(content):
        if idx < len(seg_texts):
            set_para_text(p, seg_texts[idx])

# ── FASES ──────────────────────────────────────────────────────────────────
fase1_text = 'Empecé definiendo los objetivos basándome en el temario de ASIX e investigué las vulnerabilidades más críticas de OWASP. La decisión más importante fue el sistema de doble modo: en lugar de dos proyectos separados, usé la variable APP_MODE para que el mismo código se comporte de forma diferente según la configuración, sin duplicar nada.'
fase2_text = 'Dibujé el esquema E-R en papel antes de pasarlo a Prisma y tuve que rehacerlo dos veces. El primer diseño conectaba tickets a servicios, pero era más correcto conectarlos a clientes con técnicos asignados. También el healthcheck de Docker fue un problema: Prisma intentaba conectarse antes de que PostgreSQL terminara de arrancar, así que añadí sleep 15 en el script de inicio.'
fase3_text = 'Desarrollé el servidor con Node.js y TypeScript, programando el sistema de doble modo (APP_MODE=vulnerable/secure). El mayor reto fue TypeScript: nunca lo había usado de forma seria. El compilador interrumpía el build por noImplicitAny, ajusté el tsconfig.json. También el error de los retornos de carro DOS (\\r\\n) en los scripts que Alpine leía como "illegal option -\\r" —solución: convertir a Unix \\n antes del build.'
fase4_text = 'Usé lo aprendido en Implantación de Aplicaciones Web para desarrollar la interfaz en PHP con estilo Cyberpunk/Glassmorphism. El problema más curioso fue la integración PHP-Express: el navegador enviaba cookies de sesión PHP pero el backend esperaba JWT. Tuve que hacer que PHP extrajera el token de $_SESSION y lo incluyera en cada fetch al backend.'

s, e = find_section_range(paras, 'fase 1')
if not s: s, e = find_section_range(paras, 'planificaci')
if s:
    content = get_content_paras(paras, s, e)
    if content: set_para_text(content[0][1], fase1_text)

s, e = find_section_range(paras, 'fase 2')
if not s: s, e = find_section_range(paras, 'disseny')
if s:
    content = get_content_paras(paras, s, e)
    if content: set_para_text(content[0][1], fase2_text)

s, e = find_section_range(paras, 'fase 3')
if not s: s, e = find_section_range(paras, 'backend')
if s:
    content = get_content_paras(paras, s, e)
    if content: set_para_text(content[0][1], fase3_text)

s, e = find_section_range(paras, 'fase 4')
if not s: s, e = find_section_range(paras, 'frontend corporatiu')
if s:
    content = get_content_paras(paras, s, e)
    if content: set_para_text(content[0][1], fase4_text)

# ── AMPLIACIONES ───────────────────────────────────────────────────────────
amp1 = 'El sistema ya permite simular SQLi, XSS, fuerza bruta, DoS e IDOR. La ampliación sería completar la matriz OWASP Top 10 añadiendo SSRF (Server-Side Request Forgery), LFI (Local File Inclusion) avanzado y vulnerabilidades de deserialización insegura, cada una con su contramedida activa en modo seguro. Coste estimado: 40h / 1.000€.'
amp2 = 'El monitor SOC actual usa Chart.js con polling cada 8s. La ampliación añadiría Loki para centralizar logs de todos los contenedores Docker y MaxMind GeoIP para mostrar un mapa geográfico en tiempo real de dónde provienen los ataques. Coste estimado: 24h / 600€.'
amp3 = 'Ya existe el script sofia-audit-framework.py (Python interactivo en terminal). La ampliación integraría todos sus módulos (SQLi, Brute Force, IDOR, DoS) directamente en la consola web /consola, con log detallado de cada payload enviado y respuesta recibida. Permitiría también exportar el informe de auditoría en PDF. Coste estimado: 20h / 500€.'

s, e = find_section_range(paras, '6.1')
if s:
    content = get_content_paras(paras, s, e)
    if content: set_para_text(content[0][1], amp1)

s, e = find_section_range(paras, '6.2')
if s:
    content = get_content_paras(paras, s, e)
    if content: set_para_text(content[0][1], amp2)

s, e = find_section_range(paras, '6.3')
if s:
    content = get_content_paras(paras, s, e)
    if content: set_para_text(content[0][1], amp3)

# ── CONCLUSIONES ───────────────────────────────────────────────────────────
conc_texts = [
    'Cuando empecé este proyecto tenía una idea bastante vaga: algo de ciberseguridad con Docker. Mirando el resultado final, creo que ha quedado bastante más completo y ambicioso de lo que esperaba, aunque también me ha costado mucho más trabajo del que había calculado.',
    'Lo que más me ha aportado ha sido enfrentarme a problemas reales de arquitectura: cómo hacer que varios servicios se comuniquen entre sí, cómo gestionar la autenticación de forma segura, cómo estructurar el código para que sea mantenible. Todo eso se aprende de verdad cuando lo tienes que hacer tú solo, no cuando lo ves en un tutorial.',
    'Ha habido momentos difíciles. El más duro fue cuando me di cuenta de que el diseño inicial de la BBDD no era correcto y tuve que rehacerlo a mitad del backend. También el error de los \\r\\n de Windows en los scripts —ese tipo de bug que no ves a primera vista y que te tiene horas mirando logs sin entender qué pasa.',
    'TypeScript ha sido el mayor descubrimiento. Al principio me parecía trabajo extra innecesario, pero luego vi que me ahorraba horas de debugging porque el compilador avisaba de errores que en JavaScript puro solo habrías encontrado en tiempo de ejecución. Sofia Solutions no es un proyecto perfecto, tiene cosas que mejorar, pero es un sistema funcional y real que demuestra cosas interesantes sobre la ciberseguridad en entornos empresariales.',
]
s, e = find_section_range(paras, 'conclusi')
if s:
    content = get_content_paras(paras, s, e)
    for idx, (pi, p) in enumerate(content):
        if idx < len(conc_texts):
            set_para_text(p, conc_texts[idx])

doc.save(OUT)
print(f'OK - Guardado en: {OUT}')
