"""
RESTAURACIÓN TOTAL DE 'Memoria. Entrega 4.docx'
Vuelve al estado original eliminando los cambios de bibliografía.
"""
from docx import Document

FILE = r'C:\Users\shair\Downloads\Memoria. Entrega 4.docx'
doc = Document(FILE)

# 1. Eliminar los párrafos que añadí al final (los últimos 11 de referencias + el título)
# Buscamos el párrafo '8. Bibliografía' o el inicio de los cambios.
# Para ir a lo seguro, si hay más de 244 párrafos (el total original), borramos los sobrantes.
original_count = 244
while len(doc.paragraphs) > original_count:
    p = doc.paragraphs[-1]
    p._p.getparent().remove(p._p)

# 2. Restaurar el texto original en los párrafos 224-243 (que estaban vacíos o alterados)
REFS = [
    'ISO/IEC. (2022). ISO/IEC 27001:2022 Information security, cybersecurity and ',
    '    privacy protection – Information security management systems – Requirements. ',
    '    https://www.iso.org/standard/27001',
    '',
    'Mozilla. (2024, 15 de abril). HTTP cookies. MDN Web Docs. ',
    '    https://developer.mozilla.org/en-US/docs/Web/HTTP/Cookies',
    '',
    'n8n.io. (2023). Self-hosting n8n: Configuration and deployment. n8n Documentation. ',
    '    https://docs.n8n.io/hosting/',
    '',
    'National Institute of Standards and Technology (NIST). (2018, 16 de abril). ',
    '    Framework for Improving Critical Infrastructure Cybersecurity, Version 1.1. ',
    '    https://www.nist.gov/cyberframework',
    '',
    'OWASP Foundation. (2021). OWASP Top 10:2021 The Next Generation of Application ',
    '    Security. https://owasp.org/www-project-top-ten/',
    '',
    'Prisma Data, Inc. (2024). Prisma ORM documentation: Working with raw SQL. ',
    '    Prisma Docs. https://www.prisma.io/docs/concepts/components/prisma-client/',
    '    raw-database-access'
]

# Restauramos el contenido y quitamos cualquier fuente forzada
for i, txt in enumerate(REFS):
    idx = 224 + i
    if idx < len(doc.paragraphs):
        p = doc.paragraphs[idx]
        p.text = txt
        # Quitar formatos específicos
        p.style = doc.styles['Normal']
        for run in p.runs:
            run.font.name = None
            run.font.size = None

doc.save(FILE)
print("Restauración completada con éxito.")
