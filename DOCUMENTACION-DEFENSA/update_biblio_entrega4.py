"""
Script para 'Memoria. Entrega 4.docx' — Actualiza SOLO la bibliografía.
Usa fuente Ubuntu, tamaño 12, interlineado 1.5, justificado.
"""
import copy
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

FILE = r'C:\Users\shair\Downloads\Memoria. Entrega 4.docx'
doc = Document(FILE)
paras = doc.paragraphs

# 1. Encontrar la sección 7 para saber dónde termina
idx_conclusiones = -1
for i, p in enumerate(paras):
    if '7. Conclusiones' in p.text and (p.style.name.startswith('Heading') if p.style else True):
        idx_conclusiones = i

if idx_conclusiones == -1:
    # Si no la encontramos por estilo, buscamos por texto (la última aparición)
    for i, p in enumerate(paras):
        if '7. Conclusiones' in p.text:
            idx_conclusiones = i

# 2. Encontrar el inicio de la Bibliografía (sección 8)
# Si no existe el título, lo crearemos después de las conclusiones.
idx_biblio = -1
for i in range(idx_conclusiones + 1, len(paras)):
    if 'Bibliograf' in paras[i].text:
        idx_biblio = i
        break

# Si no hay título de bibliografía, lo buscamos por el primer enlace o referencia
if idx_biblio == -1:
    for i in range(idx_conclusiones + 1, len(paras)):
        if 'http' in paras[i].text or 'Docker' in paras[i].text:
            idx_biblio = i
            break

# 3. Definir las nuevas referencias
BIBLIO_TEXTS = [
    'Docker Documentation. (2024). Docker Compose Overview. https://docs.docker.com/compose/',
    'Express.js. (2024). Express - Node.js web application framework. https://expressjs.com/',
    'Grafana Labs. (2024). Grafana documentation. https://grafana.com/docs/',
    'Microsoft. (2024). GitHub Codespaces Documentation. https://docs.github.com/en/codespaces',
    'n8n.io. (2024). n8n documentation. https://docs.n8n.io/',
    'OWASP Foundation. (2021). OWASP Top 10:2021. https://owasp.org/www-project-top-ten/',
    'PHP Documentation Group. (2024). PHP Manual. https://www.php.net/docs.php',
    'PostgreSQL Global Development Group. (2024). PostgreSQL 15 Documentation. https://www.postgresql.org/docs/15/',
    'Prisma ORM. (2024). Prisma Documentation. https://www.prisma.io/docs',
    'Prometheus Authors. (2024). Prometheus Overview. https://prometheus.io/docs/introduction/overview/',
    'TypeScript Team. (2024). TypeScript Documentation. https://www.typescriptlang.org/docs/'
]

# 4. Eliminar todo desde idx_biblio hasta el final para reemplazarlo
# (Excepto si idx_biblio es -1, en cuyo caso añadimos al final)
start_replacing = idx_biblio if idx_biblio != -1 else len(paras)

# Añadimos el título si no estaba claro
if idx_biblio == -1 or '8.' not in paras[idx_biblio].text:
    new_h = doc.add_paragraph('8. Bibliografía')
    new_h.style = 'Heading 1' # Intentamos mantener el estilo de títulos
else:
    paras[idx_biblio].text = '8. Bibliografía'

# Añadimos las referencias
for ref in BIBLIO_TEXTS:
    p = doc.add_paragraph(ref)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = 'Ubuntu'
        run.font.size = Pt(12)

# Borramos los párrafos antiguos que quedaron entre el título nuevo y el final
if idx_biblio != -1:
    # Este proceso es delicado, borramos de atrás hacia adelante
    for i in range(len(paras) - 1, idx_biblio, -1):
        # Si el párrafo ya es uno de los nuevos (añadidos al final de la lista del doc), no lo borramos
        # En python-docx, add_paragraph añade al final del body.
        pass
    # Para ser más limpios, simplemente borramos el contenido de los párrafos antiguos
    for i in range(idx_biblio + 1, start_replacing + (len(paras) - start_replacing - len(BIBLIO_TEXTS) - 1)):
        # doc.paragraphs es dinámico, mejor no iterar borrando
        pass

# Simplificación: Borramos todo el texto de los párrafos antiguos de la sección 8
if idx_biblio != -1:
    for i in range(idx_biblio + 1, len(doc.paragraphs) - len(BIBLIO_TEXTS) - 1):
        doc.paragraphs[i].text = ""

doc.save(FILE)
print(f"OK — Bibliografía actualizada con fuente Ubuntu en: {FILE}")
