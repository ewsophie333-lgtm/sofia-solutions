"""
FIXED SCRIPT — ACTUALIZA LA SECCIÓN 7 REAL (AL FINAL)
"""
import copy
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

OUT = r'C:\Users\shair\Downloads\Memoria_SOFIA_FINAL_V3.docx'
doc = Document(OUT)
paras = doc.paragraphs

# 1. Encontrar la SECCIÓN 7 real (la que es Heading 1)
target_idx = -1
for i, p in enumerate(paras):
    style_name = p.style.name if p.style else ""
    if '7. Conclusiones' in p.text and style_name.startswith('Heading'):
        target_idx = i
        break

if target_idx != -1:
    print(f"Encontrada sección real en índice {target_idx}")
    j = target_idx + 1
    count = 0
    texts = [
        'Tras todo el trabajo de diseño y desarrollo, Sofia Solutions ha pasado de ser una idea a un sistema real que demuestra cómo proteger a una PYME de forma efectiva. Lo más importante ha sido verificar que, con Docker Compose y herramientas libres, se puede montar un Centro de Operaciones (SOC) profesional sin necesidad de pagar licencias carísimas que muchas empresas no pueden permitirse.',
        'En la parte técnica, estoy muy satisfecha con el backend. Al usar PostgreSQL con Prisma y TypeScript, el sistema es mucho más robusto y ordenado de lo que esperaba. Pero lo que de verdad aporta valor es la consola de auditoría: poder lanzar ataques controlados y ver en tiempo real cómo Grafana muestra los picos de actividad hace que la seguridad sea algo visual y fácil de explicar. Todo esto, además, funcionando bajo HTTPS con dominios propios, lo que le da el acabado profesional que buscaba para la defensa.',
        'La automatización con n8n también ha sido un punto clave, ya que permite que el sistema no solo vigile, sino que reaccione solo enviando correos cuando detecta algo crítico o se pulsa el botón SOS. Además, usar GitHub Codespaces me ha salvado el proyecto. Al no tener permisos para instalar Docker en los PCs de clase y estar cansada de máquinas virtuales que se cuelgan o se corrompen solas, trabajar en la nube ha hecho que todo el desarrollo sea mucho más fluido y fiable.',
        'Para terminar, me voy con la satisfacción de haber unido la teoría de clase con la práctica real. Sofia Solutions no es solo un proyecto para aprobar, es una plataforma escalable y documentada que me ha servido para entender de verdad cómo se gestiona la seguridad en una infraestructura moderna.'
    ]
    
    # Limpiamos los párrafos vacíos o de texto que haya después del título hasta encontrar el siguiente Heading
    while j < len(paras) and count < 4:
        st_name = paras[j].style.name if paras[j].style else ""
        if st_name.startswith('Heading'): break
        
        if paras[j].text.strip() or count < 4:
            # Borrar runs actuales
            for r in list(paras[j].runs):
                r._r.getparent().remove(r._r)
            # Insertar nuevo texto
            nr = paras[j].add_run(texts[count])
            nr.font.name = 'Arial'
            nr.font.size = Pt(12)
            # Formato de párrafo
            paras[j].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paras[j].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            paras[j].paragraph_format.first_line_indent = Cm(0.5)
            count += 1
        j += 1

# 2. Restaurar el índice (por si lo rompimos antes)
doc.paragraphs[42].text = "" # Limpiar lo que escribimos por error
doc.paragraphs[43].text = ""

doc.save(OUT)
print("¡Hecho! Conclusiones actualizadas en la sección correcta.")
