"""
Part 3 — Secciones 4.5-9 + Annex
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

F = r'C:\Users\shair\Desktop\sofia-solutions\sofia-solutions\Memoria_v2.0_Sofia.docx'
doc = Document(F)

def body(doc,txt,indent=True):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=Pt(18)
    if indent: p.paragraph_format.first_line_indent=Cm(0.5)
    r=p.add_run(txt); r.font.name='Verdana'; r.font.size=Pt(12); return p

def h1(doc,txt):
    p=doc.add_heading(txt,level=1)
    for r in p.runs: r.font.name='Verdana'; r.font.size=Pt(14); r.bold=True
    return p

def h2(doc,txt):
    p=doc.add_heading(txt,level=2)
    for r in p.runs: r.font.name='Verdana'; r.font.size=Pt(12); r.bold=True
    return p

def bul(doc,txt):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(1); p.paragraph_format.space_after=Pt(3)
    r=p.add_run('• '+txt); r.font.name='Verdana'; r.font.size=Pt(12); return p

def mk_table(doc,headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Normal Table'
    row=t.add_row()
    for i,c in enumerate(headers):
        cell=row.cells[i]; cell.paragraphs[0].clear()
        r=cell.paragraphs[0].add_run(c); r.bold=True; r.font.name='Verdana'; r.font.size=Pt(10)
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'1F3564'); cell._tc.get_or_add_tcPr().append(shd)
        r.font.color.rgb=RGBColor(255,255,255)
    for rd in rows:
        row=t.add_row()
        for i,c in enumerate(rd):
            cell=row.cells[i]; cell.paragraphs[0].clear()
            r=cell.paragraphs[0].add_run(str(c)); r.font.name='Verdana'; r.font.size=Pt(10)
    doc.add_paragraph(); return t

def code(doc,fname,snip):
    t=doc.add_table(rows=2,cols=1); t.style='Normal Table'
    h=t.rows[0].cells[0]; h.paragraphs[0].clear()
    hr=h.paragraphs[0].add_run(fname); hr.bold=True; hr.font.name='Verdana'; hr.font.size=Pt(10)
    h.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'D9E1F2'); h._tc.get_or_add_tcPr().append(shd)
    c=t.rows[1].cells[0]; c.paragraphs[0].clear()
    cr=c.paragraphs[0].add_run(snip); cr.font.name='Courier New'; cr.font.size=Pt(9)
    doc.add_paragraph()

def ph(doc,desc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'[CAPTURA: {desc}]'); r.font.name='Verdana'; r.font.size=Pt(11)
    r.font.color.rgb=RGBColor(0x88,0x88,0x88)

def cap(doc,n,desc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'Figura {n}: {desc}'); r.italic=True; r.font.name='Verdana'; r.font.size=Pt(10)

# ── 4.5 CODIFICACIÓN ───────────────────────────────────────────────────────
h2(doc,'4.5 Codificación')
body(doc,'La programación está organizada por capas. El frontend PHP consume la API REST del backend. Lo más relevante del código son los controladores de autenticación dual, el middleware de detección de ataques y los scripts de automatización.')

code(doc,'docker-compose.yml — Orquestación completa del stack',
"""services:
  backend:
    build: ./sofia-backend
    environment:
      APP_MODE: ${APP_MODE:-secure}
      DATABASE_URL: postgresql://postgres:postgres@postgres:5432/sofia_solutions
      JWT_SECRET: ${JWT_SECRET:-sofia-super-secret-key-2026}
      N8N_WEBHOOK_URL: http://n8n:5678/webhook/alert
    ports: ["8001:8001"]
    depends_on: [postgres]

  postgres:
    image: postgres:15-alpine
    environment:
      POSTGRES_DB: sofia_solutions
    volumes:
      - postgres-data:/var/lib/postgresql/data

  tunnel:
    image: alpine
    command: >
      sh -c "apk add --no-cache openssh-client &&
             ssh -o StrictHostKeyChecking=no
             -R sofia-solutions:80:frontend:80 serveo.net"
""")
ph(doc,'Terminal con docker compose ps mostrando los 7 servicios en estado Up')
cap(doc,3,'Todos los servicios Docker Compose activos')

code(doc,'auth.controller.ts — Login dual V1 (vulnerable) vs V2 (seguro)',
"""// LOGIN V1 — VULNERABLE (solo para demo educativa)
export async function loginV1(req: Request, res: Response) {
  const { email, password } = req.body;
  // Query sin parametrizar — permite SQLi bypass
  const users: any[] = await prisma.$queryRawUnsafe(
    `SELECT * FROM "User" WHERE "email" = '${email}' AND "passwordHash" IS NOT NULL`
  );
  const isInjection = email.includes("'") || email.includes("--");
  if (!users[0] || (!isInjection && users[0].password !== password)) {
    throw new ApiError(401, "Invalid credentials (Vulnerable Mode)");
  }
  const token = signAccessToken({ sub: users[0].id, role: users[0].role }, "vulnerable");
  res.json({ accessToken: token });
}

// LOGIN V2 — SEGURO
export async function loginV2(req: Request, res: Response) {
  const { email, password } = req.body;
  // Consulta parametrizada con Prisma — inmune a SQLi
  const user = await prisma.user.findUnique({ where: { email } });
  if (!user || !(await bcrypt.compare(password, user.passwordHash))) {
    throw new ApiError(401, "Invalid identity credentials");
  }
  const token = signAccessToken({ sub: user.id, role: user.role }, "secure");
  res.json({ accessToken: token });
}
""")

code(doc,'attackDetection.ts — Middleware WAF con bypass para consola de auditoría',
"""export async function attackDetection(req, res, next) {
  // Bypass para la consola de auditoría (ataques controlados)
  const referer = req.headers.referer || '';
  if (referer.includes('/consola')) { return next(); }

  const attack = detectAttackPatterns({
    path: req.path, query: req.query, body: req.body, headers: req.headers
  });

  if (attack.detected) {
    await prisma.securityEvent.create({ data: {
      type: attack.type, ipOrigin: req.ip,
      payload: JSON.stringify(attack.payload), blocked: true
    }});
    return res.status(403).json({ error: 'Attack detected and blocked', type: attack.type });
  }
  next();
}
""")

code(doc,'rateLimiter.ts — Rate limiting con alerta automática a n8n',
"""export const loginRateLimiter = rateLimit({
  windowMs: 15 * 60 * 1000,  // 15 minutos
  max: 5,                     // 5 intentos por IP
  handler: (req, res) => {
    // Disparar alerta SOC via webhook n8n
    fetch(process.env.N8N_WEBHOOK_URL!, {
      method: 'POST', headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({
        tipo: 'BRUTE_FORCE_DETECTED', ip: req.ip,
        timestamp: new Date().toISOString()
      })
    }).catch(() => {});
    res.status(429).json({ error: 'Bloqueado 15 minutos.', retryAfter: '15 min' });
  }
});
""")

code(doc,'schema.prisma — Modelo de datos completo',
"""model User {
  id           Int      @id @default(autoincrement())
  email        String   @unique
  passwordHash String
  role         Role     @default(CLIENT)
  createdAt    DateTime @default(now())
  tickets      Ticket[] @relation("AssignedTickets")
}
model Client {
  id      Int      @id @default(autoincrement())
  nombre  String
  empresa String
  email   String   @unique
  plan    Plan     @default(INDIVIDUAL)
  tickets Ticket[]
}
model Ticket {
  id          Int      @id @default(autoincrement())
  titulo      String
  prioridad   Priority @default(MEDIA)
  estado      Status   @default(ABIERTO)
  clienteId   Int
  createdAt   DateTime @default(now())
  cliente     Client   @relation(fields: [clienteId], references: [id])
}
enum Role     { ADMIN CLIENT }
enum Plan     { INDIVIDUAL BUSINESS BUSINESS_MAX }
enum Priority { BAJA MEDIA ALTA CRITICA }
enum Status   { ABIERTO EN_PROCESO RESUELTO CERRADO }
""")

ph(doc,'Consola de auditoría ejecutando SQLi bypass — terminal mostrando EXPLOIT EXITOSO')
cap(doc,4,'SQL Injection bypass en modo vulnerable — Auth Bypass Clásico')
ph(doc,'Consola de auditoría — mismo ataque en modo seguro mostrando BLOQUEADO')
cap(doc,5,'SQLi bloqueado por middleware en modo seguro')
ph(doc,'Brute force completado — terminal mostrando CRACKED en posición 9/1000')
cap(doc,6,'Ataque de fuerza bruta con diccionario rockyou en modo vulnerable')

h2(doc,'4.6 Administración')
body(doc,'El sistema tiene gestión a dos niveles: infraestructura vía Docker y aplicación vía panel web.')
for cmd in [
    'Arrancar: docker compose up -d',
    'Cambiar modo: APP_MODE=vulnerable docker compose up -d',
    'Logs en tiempo real: docker compose logs -f backend',
    'Prisma Studio (BBDD visual): http://localhost:5556',
]: bul(doc,cmd)
mk_table(doc,['Funcionalidad','ADMIN','CLIENT','Sin sesión'],[
    ('Home / planes','✓','✓','✓'),
    ('Login vulnerable y seguro','✓','✓','✓'),
    ('Dashboard cliente','✓','✓','✗'),
    ('Ver/crear tickets','✓','✓','✗'),
    ('Panel SOC','✓','✗','✗'),
    ('Consola auditoría','✓','✗','✗'),
    ('Gestión de usuarios','✓','✗','✗'),
])
ph(doc,'Panel SOC con feed de eventos en tiempo real y métricas')
cap(doc,7,'Panel SOC mostrando eventos de seguridad detectados en tiempo real')
ph(doc,'Prisma Studio con tabla User mostrando contraseñas hasheadas en modo seguro')
cap(doc,8,'Prisma Studio — Base de datos con Bcrypt hashes visibles')

h2(doc,'4.7 Aspectos de seguridad')
mk_table(doc,['Vulnerabilidad OWASP','Modo VULNERABLE','Modo SEGURO'],[
    ('A03 Inyección SQL','$queryRawUnsafe con string concatenado','Prisma findUnique con parámetros tipados'),
    ('A07 Autenticación rota','Contraseña en texto plano en BD','Bcrypt 12 rounds ($2b$12$...)'),
    ('A04 Rate limiting','Sin límite → fuerza bruta libre','5 intentos/15min + bloqueo 429'),
    ('A05 XSS','Sin escape de salida','htmlspecialchars() PHP + CSP header'),
    ('IDOR','ID secuencial → enumerable','Verificación de propietario en middleware'),
])
body(doc,'HTTPS: Túnel serveo.net proporciona certificado SSL/TLS automático. Todo el tráfico exterior va cifrado. PostgreSQL no expone puerto al exterior en producción, solo es accesible desde la red interna Docker.')
doc.add_page_break()

# ── 5. FASES ───────────────────────────────────────────────────────────────
h1(doc,'5. Fases en la Realización del Proyecto')
mk_table(doc,['Fase','Tareas','Duración','Período'],[
    ('1. Planificación','Objetivos, investigación OWASP, selección tech','2 semanas','Oct 2025'),
    ('2. Diseño','Arquitectura Docker, E-R, maquetas UI','2 semanas','Nov 2025'),
    ('3. Backend (API)','Node+TS, Prisma, endpoints, modo dual','4 semanas','Nov-Dic 2025'),
    ('4. Frontend (PHP)','Interfaz Cyberpunk/Glassmorphism, panel admin','3 semanas','Ene 2026'),
    ('5. Monitorización','Prometheus, Grafana, n8n, SOC feed','2 semanas','Feb 2026'),
    ('6. Pruebas','Funcionales, ataques demo, correcciones','2 semanas','Mar 2026'),
    ('7. Documentación','Memoria, presentación, pendrive','2 semanas','Mar 2026'),
])
body(doc,'Fase 1 — Planificación: La primera decisión importante fue el sistema de doble modo. Al principio pensé en dos proyectos separados (uno vulnerable y otro seguro), pero eso era demasiado trabajo. Al final la variable de entorno APP_MODE fue la solución elegante: el mismo código se comporta de forma diferente según la configuración sin duplicar nada.')
body(doc,'Fase 2 — Diseño: Diseñé el esquema E-R con papel antes de pasarlo a Prisma y tuve que rehacerlo dos veces. El primer diseño conectaba los tickets directamente a los servicios, pero luego vi que tenía más sentido conectarlos a los clientes. Ese cambio lo hice tarde y me costó actualizar bastantes queries. También el healthcheck de Docker fue un problema al principio: Prisma intentaba conectarse antes de que PostgreSQL hubiera terminado de arrancar.')
body(doc,'Fase 3 — Backend: La fase más larga. Nunca había trabajado con TypeScript de forma seria. Hubo un momento donde el compilador TSC interrumpía el build por noImplicitAny. La solución fue ajustar tsconfig.json. También tuve el error de los retornos de carro DOS (\\r\\n) en los scripts que el intérprete Alpine leía como "illegal option -\\r" —solución: convertir a Unix \\n antes del build.')
body(doc,'Fase 4 — Frontend: El diseño Cyberpunk/Glassmorphism con CSS y Chart.js fue lo más visual. El problema más curioso fue la integración PHP-Express: el navegador enviaba cookies de sesión PHP pero el backend esperaba JWT en la cabecera Authorization. Tuve que hacer que PHP extrajera el token JWT de $_SESSION y lo pusiera en cada fetch al backend.')
body(doc,'Fase 5 — Monitorización: Configurar Prometheus fue más manual de lo esperado: añadir endpoints /metrics en la API con prom-client. n8n para las alertas automáticas también requirió varias iteraciones para que el webhook funcionara correctamente desde dentro de la red Docker.')
body(doc,'Fase 6 — Pruebas: Pruebas con los scripts propios del repo. Verifiqué que el rate limiting bloqueaba en HTTP 429 y que Grafana mostraba el pico de intentos en tiempo real.')
ph(doc,'Grafana dashboard con métrica sofia_login_attempts_total mostrando pico de ataques')
cap(doc,9,'Grafana — Pico de intentos de login durante el ataque de fuerza bruta')
doc.add_page_break()

# ── 6. AMPLIACIONES ────────────────────────────────────────────────────────
h1(doc,'6. Ampliaciones')
h2(doc,'6.1 Matriz completa de vulnerabilidades demostrables')
body(doc,'El sistema ya permite demostrar SQLi, XSS, brute force, DoS y IDOR. La ampliación sería añadir más escenarios del OWASP Top 10 con su correspondiente contramedida en modo seguro.')
mk_table(doc,['Vulnerabilidad','Estado actual','Ampliación propuesta','Coste est.'],[
    ('SQLi (A03)','✓ Implementado','UNION-based y time-based','8h / 200€'),
    ('XSS (A03)','✓ Básico','XSS reflejado y almacenado','10h / 250€'),
    ('IDOR (A01)','✓ Tickets SOC','Recursos de clientes entre sí','6h / 150€'),
    ('SSRF (A10)','✗ No implementado','Endpoint vulnerable con fetch interno','12h / 300€'),
])

h2(doc,'6.2 Monitor SOC en vivo con telemetría avanzada')
body(doc,'El SOC actual usa Chart.js con polling cada 8s. La ampliación sería integrar Loki para centralizar logs de todos los contenedores y añadir un mapa geográfico con MaxMind GeoIP mostrando de dónde vienen los ataques. Coste estimado: 24h / 600€ adicionales.')

h2(doc,'6.3 Herramienta de auditoría — Sofia Exploit Kit')
body(doc,'Ya existe scripts/sofia-audit-framework.py (Python, interactivo). La ampliación sería integrarlo en la consola web /consola de forma que los módulos sean ejecutables desde el navegador con log detallado de comandos y payloads reales.')
ph(doc,'Sofia Audit Framework Python ejecutándose en terminal con menú interactivo')
cap(doc,10,'Sofia Audit Framework — Herramienta de auditoría CLI en terminal')

h2(doc,'6.4 Aplicación móvil para técnicos')
body(doc,'La API REST ya está documentada. Con React Native se podría desarrollar una app para iOS/Android que permita a los técnicos consultar y actualizar tickets asignados desde el teléfono. Estimación: 80h / 2.000€.')

mk_table(doc,['Ampliación','Horas','Coste','Prioridad'],[
    ('6.1 Más vulnerabilidades demo','40h','1.000€','Alta'),
    ('6.2 SOC con Loki + GeoIP','24h','600€','Media'),
    ('6.3 Consola web exploit kit','20h','500€','Media'),
    ('6.4 App móvil técnicos','80h','2.000€','Baja'),
    ('TOTAL','164h','4.100€','—'),
])
doc.add_page_break()

# ── 7. CONCLUSIONES ────────────────────────────────────────────────────────
h1(doc,'7. Conclusiones')
body(doc,'Cuando empecé este proyecto tenía una idea bastante vaga: algo de ciberseguridad con Docker. Mirando el resultado final, creo que ha quedado bastante más completo y ambicioso de lo que esperaba, aunque también me ha costado mucho más trabajo del que había calculado.')
body(doc,'Lo que más me ha aportado ha sido enfrentarme a problemas reales de arquitectura: cómo hacer que varios servicios se comuniquen entre sí, cómo gestionar la autenticación de forma segura, cómo estructurar el código para que sea mantenible. Todo eso se aprende de verdad cuando lo tienes que hacer tú solo, no cuando lo ves en un tutorial.')
body(doc,'Ha habido momentos difíciles. El más duro fue cuando me di cuenta de que el diseño inicial de la BBDD no era correcto y tuve que rehacerlo a mitad del backend. También el error de los \\r\\n de Windows en los scripts —ese tipo de bug que no ves a primera vista y que te tiene horas mirando logs sin entender qué pasa.')
body(doc,'Por otro lado, hay cosas que han salido mejor de lo esperado. El sistema de doble modo quedó muy limpio con solo la variable APP_MODE. El panel SOC con Chart.js y el feed actualizado cada 8 segundos quedó bastante profesional. Y Grafana, una vez configurado con las métricas de prom-client, da mucho juego para la demo ante el tribunal.')
body(doc,'TypeScript ha sido el mayor descubrimiento. Al principio me parecía trabajo extra innecesario, pero luego vi que me ahorraba horas de debugging porque el compilador avisaba de errores que en JavaScript puro solo habrías encontrado en tiempo de ejecución.')
body(doc,'Sofia Solutions no es un proyecto perfecto, tiene cosas que mejorar, pero es un sistema funcional y real que demuestra cosas interesantes sobre la ciberseguridad en entornos empresariales. Y eso era exactamente lo que quería conseguir.')
doc.add_page_break()

# ── 8. BIBLIOGRAFÍA ────────────────────────────────────────────────────────
h1(doc,'8. Bibliografía')
refs = [
    'Docker Documentation. (2024). Compose file reference. https://docs.docker.com/compose/compose-file/',
    'ISO/IEC. (2022). ISO/IEC 27001:2022 — Information security management systems. https://www.iso.org/standard/27001',
    'Mozilla Developer Network. (2024). HTTP cookies. https://developer.mozilla.org/docs/Web/HTTP/Cookies',
    'n8n.io. (2023). Self-hosting n8n: Configuration and deployment. https://docs.n8n.io/hosting/',
    'NIST. (2018). Framework for Improving Critical Infrastructure Cybersecurity v1.1. https://www.nist.gov/cyberframework',
    'OWASP Foundation. (2021). OWASP Top 10:2021. https://owasp.org/www-project-top-ten/',
    'Prisma Data. (2024). Prisma ORM documentation. https://www.prisma.io/docs/',
    'Prometheus Authors. (2024). Prometheus documentation. https://prometheus.io/docs/',
    'Grafana Labs. (2024). Grafana documentation. https://grafana.com/docs/',
    'Ali Syed, B. (2023). TypeScript Deep Dive. https://basarat.gitbook.io/typescript/',
    'PostgreSQL Global Development Group. (2024). PostgreSQL 15 Documentation. https://www.postgresql.org/docs/15/',
    'Express.js. (2024). Express 4.x API Reference. https://expressjs.com/en/api.html',
]
for ref in refs:
    p=doc.add_paragraph()
    p.paragraph_format.left_indent=Cm(1.27)
    p.paragraph_format.first_line_indent=Cm(-1.27)
    r=p.add_run(ref); r.font.name='Verdana'; r.font.size=Pt(11)
doc.add_page_break()

# ── 9. SOPORTE FÍSICO ──────────────────────────────────────────────────────
h1(doc,'9. PFC en Suport Físic (Pendrive USB)')
body(doc,'Junto a la memoria impresa se entrega un pendrive USB con el contenido íntegro del proyecto.')
mk_table(doc,['Carpeta / Archivo','Descripción'],[
    ('/sofia-solutions/','Directorio raíz del proyecto completo'),
    ('  /sofia-backend/','Código fuente del servidor Node.js + TypeScript'),
    ('  /frontend-php/','Interfaz web PHP (portal corporativo y SOC)'),
    ('  /scripts/','Scripts de auditoría y automatización'),
    ('  /DOCUMENTACION-DEFENSA/','Guiones de ataque, guía de defensa'),
    ('  docker-compose.yml','Orquestación de todos los servicios'),
    ('  .devcontainer/','Configuración GitHub Codespaces'),
    ('  n8n_workflow_sofia.json','Flujo de trabajo de alertas SOS (n8n)'),
    ('/MEMORIA_TFG.docx','Documento de memoria del proyecto'),
    ('/PRESENTACION_SOFIA.html','Presentación interactiva del proyecto'),
    ('/README.md','Instrucciones de despliegue y uso'),
])
body(doc,'Despliegue desde cero: sh ./scripts/SOFIA-INICIAR.sh → acceder a http://localhost:8000')
doc.add_page_break()

# ── ANNEX I ────────────────────────────────────────────────────────────────
h1(doc,"Annex I: Pla d'Empresa")
body(doc,'Sofia Solutions nace como respuesta a una necesidad real: las PYMEs carecen de recursos para contratar un equipo de ciberseguridad, pero están igualmente expuestas. El modelo de negocio es SaaS (Software as a Service) con facturación mensual por niveles.')
mk_table(doc,['Plan','Precio/mes','Servicios incluidos'],[
    ('Essential','299 €/mes','Monitorización 8x5, alertas email, 1 auditoría trimestral'),
    ('Business','699 €/mes','Monitorización 24x7, SOC dashboard, 2 auditorías/año'),
    ('Enterprise','1.499 €/mes','SOC dedicado 24x7, IR Retainer, auditorías ilimitadas, SLA 15 min'),
])
body(doc,'El público objetivo son PYMEs de entre 10 y 200 empleados en sectores regulados como finanzas, seguros o sanidad. Estos sectores están sujetos a normativas como ISO 27001, ENS o GDPR que obligan a implementar controles de seguridad auditables.')
mk_table(doc,['Concepto','Importe anual'],[
    ('Ingresos estimados (5 clientes Essential + 2 Business)','26.580 €'),
    ('Costes de infraestructura (VPS, dominio, certificados)','1.200 €'),
    ('Costes de herramientas (n8n Cloud, Grafana Cloud)','960 €'),
    ('Costes de marketing y presencia web','600 €'),
    ('Margen bruto estimado','23.820 €'),
])
body(doc,'La viabilidad económica es clara dado que la inversión inicial es prácticamente nula al usar herramientas de código abierto y contenedores Docker como infraestructura de despliegue.')

doc.save(F)
print(f"COMPLETO. Documento guardado en:\n{F}")
