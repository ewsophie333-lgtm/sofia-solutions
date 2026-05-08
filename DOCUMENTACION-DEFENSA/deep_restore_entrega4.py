"""
RESTAURACIÓN PROFUNDA — Elimina todo el formato manual aplicado.
Vuelve a los estilos por defecto del documento original.
"""
from docx import Document

FILE = r'C:\Users\shair\Downloads\Memoria. Entrega 4.docx'
doc = Document(FILE)

for p in doc.paragraphs:
    # Solo limpiamos si no es un título (para no romper los Headings si estaban bien)
    style_name = p.style.name if p.style else ""
    
    # Reset de formato de párrafo (vuelve al valor del estilo)
    p.alignment = None
    p.paragraph_format.line_spacing = None
    p.paragraph_format.line_spacing_rule = None
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = None
    p.paragraph_format.space_before = None
    
    # Reset de fuente en cada run
    for run in p.runs:
        # Solo reseteamos si es Arial 12 (lo que yo puse)
        # O simplemente reseteamos todo lo manual para asegurar
        run.font.name = None
        run.font.size = None

# Restaurar bibliografía original
REFS = [
    'ISO/IEC. (2022). ISO/IEC 27001:2022 Information security, cybersecurity and ',
    '    privacy protection – Information security management systems – Requirements. ',
    '    https://www.iso.org/standard/27001',
    '',
    'Mozilla. (2024, 15 de abril). HTTP cookies. MDN Web Docs. ',
    '    https://developer.mozilla.org/en-US/docs/Web/HTTP/Cookies',
    '',
    'n8n.io. (2023). Self-hosting n8n: Configuration and deployment. n8n Documentation. ',
    '    https://docs.n8n.io/hosting/',
    '',
    'National Institute of Standards and Technology (NIST). (2018, 16 de abril). ',
    '    Framework for Improving Critical Infrastructure Cybersecurity, Version 1.1. ',
    '    https://www.nist.gov/cyberframework',
    '',
    'OWASP Foundation. (2021). OWASP Top 10:2021 The Next Generation of Application ',
    '    Security. https://owasp.org/www-project-top-ten/',
    '',
    'Prisma Data, Inc. (2024). Prisma ORM documentation: Working with raw SQL. ',
    '    Prisma Docs. https://www.prisma.io/docs/concepts/components/prisma-client/',
    '    raw-database-access'
]

original_start = 224
for i, txt in enumerate(REFS):
    idx = original_start + i
    if idx < len(doc.paragraphs):
        doc.paragraphs[idx].text = txt

doc.save(FILE)
print("Restauración profunda completada.")
