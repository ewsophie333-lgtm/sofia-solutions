"""
Patch Final v2.0 (FIKED) — Simplifica el lenguaje técnico, añade n8n y Codespaces.
"""
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'C:\Users\shair\Downloads\Memoria_v2.0_Sofia.docx'
doc = Document(OUT)
paras = doc.paragraphs

def set_text(p, txt):
    rpr = None
    if p.runs and p.runs[0]._r.find(qn('w:rPr')) is not None:
        rpr = copy.deepcopy(p.runs[0]._r.find(qn('w:rPr')))
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    nr = p.add_run(txt)
    if rpr is not None:
        nr._r.insert(0, rpr)

def insert_placeholder_after_oxml(para_element, label):
    """Inserta un párrafo de placeholder después de un elemento OXML."""
    new_p = OxmlElement('w:p')
    # Copiar propiedades de párrafo del original
    orig_pPr = para_element.find(qn('w:pPr'))
    if orig_pPr is not None:
        new_p.append(copy.deepcopy(orig_pPr))
    r = OxmlElement('w:r')
    t = OxmlElement('w:t'); t.text = f'(Insertar aquí captura: {label})'
    r.append(t)
    new_p.append(r)
    para_element.addnext(new_p)
    return new_p

# ── 1. CONCLUSIONES — Simplificar lenguaje ───────────────────────────
p262 = paras[262]
if 'La integración de la telemetría' in p262.text or 'Poder ver en tiempo real' in p262.text:
    set_text(p262, 
        'Poder ver en tiempo real lo que pasa en el sistema gracias a las gráficas de Grafana y Prometheus ha sido clave. '
        'Ya no son solo datos aburridos, sino que ves el ataque ocurriendo en directo. Además, n8n hace que el sistema '
        'responda solo: si alguien intenta entrar a la fuerza o se pulsa el botón SOS, llega un correo al momento.'
    )

p263 = paras[263]
if 'Finalmente, Sofia Solutions' in p263.text or 'Trabajar con GitHub Codespaces' in p263.text:
    set_text(p263,
        'Trabajar con GitHub Codespaces me ha quitado el dolor de cabeza de configurar el PC; abres el navegador y ya '
        'tienes todo listo para funcionar. Gracias a que usamos Docker, el proyecto puede desplegarse en segundos en '
        'cualquier sitio y sé que funcionará exactamente igual que en mi ordenador local.'
    )

p264 = paras[264]
set_text(p264,
    'En cuanto a las herramientas, GitHub Codespaces ha sido una de las mejores decisiones. Poder tener todo el '
    'entorno configurado en la nube me ha permitido trabajar desde cualquier sitio sin preocuparme de dependencias. '
    'Grafana y Prometheus me han sorprendido mucho: al principio pensaba que serían muy complejas, pero ver cómo '
    'pueden transformar los logs en gráficas que se mueven en directo hace que la seguridad sea mucho más fácil '
    'de entender y de explicar.'
)

# ── 2. FASE 5 — Añadir n8n ───────────────────────────────────────────
for i, p in enumerate(paras):
    if 'Fase 5: Monitorización' in p.text:
        for j in range(i+1, i+10):
            if 'Grafana fue lo más visual' in paras[j].text:
                p_grafana = paras[j]
                
                # Texto explicativo n8n
                new_p = OxmlElement('w:p')
                orig_pPr = p_grafana._p.find(qn('w:pPr'))
                if orig_pPr is not None: new_p.append(copy.deepcopy(orig_pPr))
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = ('n8n ha sido la pieza que ha permitido automatizar las respuestas. He configurado un flujo '
                          'que recibe los avisos del backend y los convierte en correos electrónicos reales. De esta '
                          'forma, si ocurre algo grave como un ataque de fuerza bruta o si el usuario pulsa el botón '
                          'de emergencia (SOS), el equipo de seguridad recibe un aviso inmediato con todos los detalles.')
                r.append(t)
                new_p.append(r)
                p_grafana._p.addnext(new_p)
                
                # Placeholder captura n8n
                insert_placeholder_after_oxml(new_p, 'Flujo de n8n: Webhook -> Set -> Send Email')
                break
        break

doc.save(OUT)
print(f'OK — Guardado en {OUT}')
