"""
Part 2 — Secciones 1-4 de la Memoria
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

F = r'C:\Users\shair\Desktop\sofia-solutions\sofia-solutions\Memoria_v2.0_Sofia.docx'
doc = Document(F)

def body(doc,txt,indent=True):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=Pt(18)
    if indent: p.paragraph_format.first_line_indent=Cm(0.5)
    r=p.add_run(txt); r.font.name='Verdana'; r.font.size=Pt(12); return p

def h1(doc,txt):
    p=doc.add_heading(txt,level=1)
    for r in p.runs: r.font.name='Verdana'; r.font.size=Pt(14); r.bold=True
    return p

def h2(doc,txt):
    p=doc.add_heading(txt,level=2)
    for r in p.runs: r.font.name='Verdana'; r.font.size=Pt(12); r.bold=True
    return p

def bul(doc,txt):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(1); p.paragraph_format.space_after=Pt(3)
    r=p.add_run('• '+txt); r.font.name='Verdana'; r.font.size=Pt(12); return p

def mk_table(doc,headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Normal Table'
    row=t.add_row()
    for i,c in enumerate(headers):
        cell=row.cells[i]; cell.paragraphs[0].clear()
        r=cell.paragraphs[0].add_run(c); r.bold=True; r.font.name='Verdana'; r.font.size=Pt(10)
        shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'1F3564'); cell._tc.get_or_add_tcPr().append(shd)
        r.font.color.rgb=RGBColor(255,255,255)
    for rd in rows:
        row=t.add_row()
        for i,c in enumerate(rd):
            cell=row.cells[i]; cell.paragraphs[0].clear()
            r=cell.paragraphs[0].add_run(str(c)); r.font.name='Verdana'; r.font.size=Pt(10)
    doc.add_paragraph(); return t

def code(doc,fname,snip):
    t=doc.add_table(rows=2,cols=1); t.style='Normal Table'
    h=t.rows[0].cells[0]; h.paragraphs[0].clear()
    hr=h.paragraphs[0].add_run(fname); hr.bold=True; hr.font.name='Verdana'; hr.font.size=Pt(10)
    h.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'D9E1F2'); h._tc.get_or_add_tcPr().append(shd)
    c=t.rows[1].cells[0]; c.paragraphs[0].clear()
    cr=c.paragraphs[0].add_run(snip); cr.font.name='Courier New'; cr.font.size=Pt(9)
    doc.add_paragraph()

def ph(doc,desc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'[CAPTURA: {desc}]'); r.font.name='Verdana'; r.font.size=Pt(11)
    r.font.color.rgb=RGBColor(0x88,0x88,0x88)

def cap(doc,n,desc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'Figura {n}: {desc}'); r.italic=True; r.font.name='Verdana'; r.font.size=Pt(10)

# ── 1. INTRODUCCIÓN ────────────────────────────────────────────────────────
h1(doc,'1. Introducción')
body(doc,'Al fijarnos en el panorama empresarial de hoy, se ve claro que muchísimas PYMEs tienen un problema de seguridad enorme porque, sencillamente, no pueden permitirse grandes inversiones en infraestructura de ciberseguridad. Contratar un SOC tradicional puede costar decenas de miles de euros al año, algo inasumible para una empresa pequeña. Por eso quise centrar mi proyecto en algo útil de verdad: montar un entorno completo que simule lo que necesitaría una PYME real para gestionar su seguridad IT, usando exclusivamente software libre y contenedores Docker.')
body(doc,'La plataforma Sofia Solutions es el resultado de ese trabajo. Es un sistema web corporativo con dos partes: un frontend en PHP con panel de administración para gestionar clientes, servicios y tickets; y un backend en Node.js con TypeScript que incorpora una API capaz de funcionar en dos modos —vulnerable y seguro— para hacer demostraciones educativas de ataques reales y sus contramedidas. Todo se complementa con Grafana y Prometheus para monitorizar el sistema en tiempo real.')
body(doc,'Módulos del ciclo aplicados:',indent=False)
for m in [
    'Seguridad y Alta Disponibilidad → base para entender vectores OWASP y defensas',
    'Implantación de Aplicaciones Web → frontend PHP, HTML, CSS',
    'Administración de Sistemas Operativos → scripts Bash/PowerShell de automatización',
    'Servicios de Red e Internet → Docker, exposición de puertos, túnel SSL (serveo.net)',
    'Administración de SGBD → diseño E-R PostgreSQL, Prisma ORM',
]: bul(doc,m)
body(doc,'El proyecto ha evolucionado bastante desde la idea inicial. Al principio solo quería montar un entorno vulnerable para practicar, pero se fue convirtiendo en algo más ambicioso: plataforma completa con panel SOC, gestión de clientes, sistema de tickets y monitorización real.')
doc.add_page_break()

# ── 2. OBJETIVOS ────────────────────────────────────────────────────────────
h1(doc,'2. Objetivos')
body(doc,'Objetivo principal: montar un laboratorio de ciberseguridad funcional que sirva como plataforma de demostración para una PYME, comparando una infraestructura sin protecciones frente a una correctamente asegurada.')
body(doc,'Objetivos específicos:',indent=False)
for o in [
    'Diseñar e implementar arquitectura de microservicios con Docker Compose (frontend, backend, BBDD, monitorización aislados).',
    'Crear API REST en Node.js+TypeScript con sistema de doble modo (APP_MODE=secure/vulnerable) para demos educativas de OWASP Top 10.',
    'Desarrollar frontend corporativo en PHP con panel admin para gestionar clientes, servicios IT y tickets de incidencias.',
    'Implementar panel SOC en tiempo real con Grafana+Prometheus y feed de actividad actualizado cada 8 segundos.',
    'Configurar acceso HTTPS público mediante túnel serveo.net sin necesidad de IP fija ni VPS contratado.',
    'Automatizar despliegue completo con scripts Bash/PowerShell (un solo comando levanta todo).',
    'Aplicar medidas de seguridad reales en modo seguro: Bcrypt (12 rounds), JWT, CSP, rate limiting, consultas parametrizadas.',
    'Documentar todo para que cualquier administrador pueda replicar el entorno.',
]: bul(doc,o)
doc.add_page_break()

# ── 3. ANÁLISIS DE REQUERIMIENTOS ──────────────────────────────────────────
h1(doc,'3. Análisis de Requerimientos')
h2(doc,'3.1 Necesidades y situación actual')
body(doc,'Cliente: Sofia Solutions, PYME de servicios IT. Situación inicial: gestión con hojas de cálculo y correos, sin sistema centralizado ni trazabilidad de incidentes. Partíamos de cero, sin ningún servidor configurado.')
body(doc,'Requerimientos funcionales identificados:',indent=False)
for r in [
    'Panel web para gestión de clientes y servicios IT',
    'Sistema de tickets para incidencias de seguridad con priorización (BAJA/MEDIA/ALTA/CRITICA)',
    'API de autenticación con roles ADMIN / USUARIO',
    'Panel SOC con feed de actividad en tiempo real (actualización cada 8s)',
    'Módulo de demo educativa de vulnerabilidades (modo dual)',
    '3 planes de servicio: Individual 499€/mes, Business 1.500€/mes, Business Max 4.200€/mes',
]: bul(doc,r)
body(doc,'Requerimientos de seguridad (modo seguro):',indent=False)
for r in [
    'JWT con expiración 8h',
    'Bcrypt con 12 rounds para contraseñas',
    'Rate limiting: 5 intentos/IP en 15 minutos',
    'Consultas parametrizadas con Prisma (sin SQLi)',
]: bul(doc,r)

h2(doc,'3.2 Estudio de alternativas tecnológicas')
mk_table(doc,['Opción','Motivo descarte'],[
    ('WordPress + plugins','Sin control fino de API, difícil dockerizar, no permite modo dual'),
    ('Django (Python)','Curva alta, menos experiencia personal'),
    ('Laravel (PHP)','Más pesado, menos flexibilidad para sistema de doble modo'),
    ('Node.js + Express + TypeScript ✓','Ligero, tipado fuerte, flexible, permite APP_MODE'),
])
mk_table(doc,['Base de datos','Motivo descarte'],[
    ('MySQL','Licencia dual, menos tipos avanzados'),
    ('SQLite','No apta para múltiples conexiones simultáneas'),
    ('MongoDB','No relacional, no apropiada para datos cliente/ticket'),
    ('PostgreSQL 15 ✓','Robusta, open source 100%, excelente integración Prisma'),
])
mk_table(doc,['Monitorización','Motivo descarte'],[
    ('Zabbix','Muy pesado, difícil dockerizar'),
    ('ELK Stack','RAM enorme, excesivo para el alcance'),
    ('Prometheus + Grafana ✓','Ligero, dockerizable, dashboards modernos'),
])

h2(doc,'3.3 Valoración económica')
mk_table(doc,['Concepto','Coste'],[
    ('VPS Hetzner CX21 (2vCPU, 4GB RAM)','5,83 €/mes'),
    ('Dominio .com','0,92 €/mes'),
    ('SSL (Let\'s Encrypt)','0 €'),
    ('Total infraestructura','7 €/mes / 84 €/año'),
])
mk_table(doc,['Fase desarrollo','Horas','Coste'],[
    ('Análisis y diseño','20h','500 €'),
    ('Backend (API)','40h','1.000 €'),
    ('Frontend (PHP)','30h','750 €'),
    ('Docker + automatización','15h','375 €'),
    ('Monitorización','10h','250 €'),
    ('Pruebas + documentación','20h','500 €'),
    ('TOTAL','135h','3.375 €'),
])
body(doc,'Comparativa: solución propietaria equivalente (Splunk SIEM + licencias) → 15.000–40.000 €/año solo en licencias.')
doc.add_page_break()

# ── 4. IMPLEMENTACIÓN ──────────────────────────────────────────────────────
h1(doc,'4. Implementación')
body(doc,'Arquitectura por microservicios con Docker Compose. Cada servicio en contenedor independiente comunicados por red interna sofia-network (bridge).')

h2(doc,'4.1 Diagramas y Esquemas')
ph(doc,'Diagrama de arquitectura Docker Compose con los 7 servicios')
cap(doc,1,'Arquitectura de microservicios Sofia Solutions con Docker Compose')
doc.add_paragraph()

mk_table(doc,['Ruta','Descripción','Acceso'],[
    ('/','Página principal corporativa','Público'),
    ('/login','Autenticación (modo vulnerable)','Público'),
    ('/login-secure','Autenticación segura (CAPTCHA + rate limiting)','Público'),
    ('/dashboard','Panel principal del cliente','Auth requerida'),
    ('/admin/security-monitor','Panel SOC completo con telemetría','Solo ADMIN'),
    ('/consola','Consola maestra de auditoría','Solo ADMIN'),
    ('POST /api/v1/auth/login','Login vulnerable (sin Bcrypt)','API pública'),
    ('POST /api/v2/auth/login','Login seguro (Bcrypt + rate limit)','API pública'),
    ('GET /metrics','Métricas Prometheus','Interno'),
])

mk_table(doc,['Servicio','Imagen','Puerto','Función'],[
    ('frontend','php:apache (custom)','8000','Web PHP corporativa'),
    ('backend','node:20-alpine (custom)','8001','API REST TypeScript'),
    ('postgres','postgres:15-alpine','5432','Base de datos'),
    ('prometheus','prom/prometheus:latest','9090','Métricas'),
    ('grafana','grafana/grafana:latest','3000','Dashboards'),
    ('n8n','n8nio/n8n:latest','5678','Alertas SOC automáticas'),
    ('tunnel','alpine (ssh)','—','Túnel SSL serveo.net'),
])

body(doc,'Diagrama E-R — Entidades principales:',indent=False)
for e in [
    'User (id, email, password[hash/plaintext según modo], role: ADMIN|CLIENT, createdAt)',
    'Client (id, nombre, empresa, email, plan: individual|business|business_max)',
    'Ticket (id, titulo, descripcion, prioridad, estado, clienteId, assignedToId, createdAt)',
    'SecurityEvent (id, tipo, ip_origen, payload, bloqueado, timestamp)',
    'Service (id, nombre, categoria, precio, descripcion)',
]: bul(doc,e)
ph(doc,'Diagrama Entidad-Relación PostgreSQL')
cap(doc,2,'Diagrama E-R de la base de datos Sofia Solutions')

h2(doc,'4.2 Herramientas de software')
mk_table(doc,['Herramienta','Versión','Uso'],[
    ('Docker Engine + Compose','v2','Orquestación de todos los servicios'),
    ('Node.js + TypeScript','18 LTS / 5.x','API REST del backend, tipado estático'),
    ('Express.js','4.x','Routing, middlewares de auth y seguridad'),
    ('Prisma ORM','5.x','Esquema BBDD, migraciones, Prisma Studio'),
    ('PostgreSQL','15-alpine','SGBD relacional, volumen persistente'),
    ('PHP + Apache','8.2 / 2.4','Frontend corporativo'),
    ('Chart.js','4.x','Gráficas de telemetría en tiempo real'),
    ('Prometheus + prom-client','latest','Métricas sofia_login_attempts_total'),
    ('Grafana','10','Dashboards de seguridad'),
    ('n8n','latest','Automatización de alertas SOS'),
])

h2(doc,'4.3 Herramientas de hardware')
body(doc,'Desarrollo en portátil personal (Intel Core i5-11ª, 16GB RAM, SSD 512GB). Consumo con todos los servicios levantados: ~4-5 GB RAM. Para producción real recomendado: 4 vCPU, 8 GB RAM, 40 GB SSD, Ubuntu 22.04 LTS.')

h2(doc,'4.4 Lenguajes')
mk_table(doc,['Lenguaje','Versión','Uso en el proyecto'],[
    ('PHP','8.2','Frontend: vistas, sesiones, lógica presentación'),
    ('HTML5 + CSS3','—','Estructura y estilos (Glassmorphism, neon Cyberpunk)'),
    ('JavaScript','ES2022','SOC feed (polling 8s), Chart.js, validaciones'),
    ('TypeScript','5.x','Backend completo: API, middlewares, controladores'),
    ('SQL','PostgreSQL 15','Schema, queries raw para demos SQLi'),
    ('YAML','—','docker-compose.yml, prometheus.yml'),
    ('Bash','5.x','scripts/sofia-iniciar.sh, sofia-ataques.sh'),
    ('PowerShell','5.1+','SOFIA-INICIAR.ps1 (para Windows)'),
    ('Prisma Schema','5.x','Modelado BBDD y migraciones'),
])

doc.save(F)
print("Part 2 OK - Secciones 1-4 guardadas")
