"""
Script final: VSCode code blocks + insertar imágenes reales en los Figura X
"""
import copy, shutil
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r'C:\Users\shair\Downloads\Memoria_v2.0_Sofia.docx'
OUT = SRC
IMG = r'C:\Users\shair\.gemini\antigravity\brain\db019609-944c-45ca-ba07-c510cd559e6c'

doc = Document(SRC)

# ── VSCode block helper ──────────────────────────────────────────────────
def vscode_block(after_para, filename, code_text):
    tbl = doc.add_table(rows=2, cols=1)
    # Quitar estilo para evitar KeyError
    tbl._tbl.attrib.pop(qn('w:styleId'), None)

    # Fila 1: barra título azul VSCode
    tc0 = tbl.rows[0].cells[0]
    tc0.paragraphs[0].clear()
    shd0 = OxmlElement('w:shd'); shd0.set(qn('w:fill'),'007ACC'); shd0.set(qn('w:val'),'clear')
    tc0._tc.get_or_add_tcPr().append(shd0)
    r0 = tc0.paragraphs[0].add_run('  ● ● ●   ' + filename)
    r0.font.name='Consolas'; r0.font.size=Pt(9); r0.font.bold=True
    r0.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)

    # Fila 2: código oscuro
    tc1 = tbl.rows[1].cells[0]
    # Eliminar párrafo por defecto
    for existing in list(tc1.paragraphs):
        existing._p.getparent().remove(existing._p)
    shd1 = OxmlElement('w:shd'); shd1.set(qn('w:fill'),'1E1E1E'); shd1.set(qn('w:val'),'clear')
    tc1._tc.get_or_add_tcPr().append(shd1)

    first = True
    for line in code_text.split('\n'):
        if first:
            p = tc1.add_paragraph(); first = False
        else:
            p = tc1.add_paragraph()
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
        p.paragraph_format.line_spacing=Pt(13)
        # fondo oscuro en párrafo
        pPr = p._p.get_or_add_pPr()
        shd_p = OxmlElement('w:shd'); shd_p.set(qn('w:fill'),'1E1E1E'); shd_p.set(qn('w:val'),'clear')
        pPr.append(shd_p)
        run = p.add_run(line if line else ' ')
        run.font.name='Consolas'; run.font.size=Pt(9)
        # Colores sintaxis básicos
        stripped = line.strip()
        if stripped.startswith('#') or stripped.startswith('//'):
            run.font.color.rgb=RGBColor(0x6A,0x99,0x55)
        elif any(stripped.startswith(k) for k in ['export','const','let','async','function','return','if','import','throw']):
            run.font.color.rgb=RGBColor(0x56,0x9C,0xD6)
        elif any(stripped.startswith(k) for k in ['services:','image:','build:','environment:','volumes:','depends','model ','enum ']):
            run.font.color.rgb=RGBColor(0x9C,0xDC,0xFE)
        elif "'" in line or '"' in line:
            run.font.color.rgb=RGBColor(0xCE,0x91,0x78)
        else:
            run.font.color.rgb=RGBColor(0xD4,0xD4,0xD4)

    # Mover tabla después del párrafo indicado
    tbl._tbl.getparent().remove(tbl._tbl)
    after_para._p.addnext(tbl._tbl)
    # Espaciado post-tabla
    sp = OxmlElement('w:p'); tbl._tbl.addnext(sp)

# ── Imagen helper ─────────────────────────────────────────────────────────
def insert_image_after(para, img_path, width_cm=14):
    """Reemplaza párrafo placeholder con imagen real."""
    try:
        run = para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = para.add_run()
        run2.add_picture(img_path, width=Cm(width_cm))
    except Exception as e:
        para.clear()
        para.add_run(f'[Imagen no disponible: {e}]')

# ── CÓDIGO MAP: keyword en Heading3 → (filename, code) ───────────────────
CODES = {
    'Docker Compose': ('docker-compose.yml', """\
services:
  backend:
    build: ./sofia-backend
    environment:
      APP_MODE: ${APP_MODE:-secure}
      DATABASE_URL: postgresql://postgres:postgres@postgres:5432/sofia_solutions
      N8N_WEBHOOK_URL: http://n8n:5678/webhook/alert
    depends_on: [postgres]
  prometheus:
    image: prom/prometheus:latest
    volumes:
      - ./prometheus.yml:/etc/prometheus/prometheus.yml
  tunnel:
    image: alpine
    command: ssh -R sofia-solutions:80:frontend:80 serveo.net"""),
    'Esquema de base de datos': ('prisma/schema.prisma', """\
model User {
  id           Int      @id @default(autoincrement())
  email        String   @unique
  passwordHash String
  role         Role     @default(CLIENT)
  createdAt    DateTime @default(now())
}
model Ticket {
  id        Int      @id @default(autoincrement())
  titulo    String
  prioridad Priority @default(MEDIA)
  clienteId Int
  cliente   Client   @relation(fields: [clienteId], references: [id])
}
enum Role     { ADMIN CLIENT }
enum Priority { BAJA MEDIA ALTA CRITICA }"""),
    'Middleware de autenticaci': ('src/middleware/auth.ts', """\
export function requireAuth(req: Request, res: Response, next: NextFunction) {
  const isAdmin = req.path.startsWith('/api/admin');
  const suffix  = isAdmin ? '_ADMIN' : '_CLIENT';
  const token   = req.cookies[`accessToken${suffix}`]
                || req.headers.authorization?.split(' ')[1];
  if (!token) throw new ApiError(401, 'Authentication required');
  req.user = verifyAccessToken(token);
  next();
}"""),
    'Sistema de doble modo': ('src/controllers/auth.controller.ts', """\
// MODO VULNERABLE — permite SQL Injection bypass
export async function loginV1(req: Request, res: Response) {
  const { email, password } = req.body;
  // Query sin parametrizar — vulnerable a SQLi
  const users: any[] = await prisma.$queryRawUnsafe(
    `SELECT * FROM "User" WHERE "email" = '${email}'`
  );
  const isInjection = email.includes("'") || email.includes("--");
  if (!users[0] || (!isInjection && users[0].password !== password))
    throw new ApiError(401, 'Invalid credentials');
  res.json({ accessToken: signAccessToken({ sub: users[0].id }) });
}

// MODO SEGURO — Bcrypt + Prisma parametrizado (inmune a SQLi)
export async function loginV2(req: Request, res: Response) {
  const { email, password } = req.body;
  const user = await prisma.user.findUnique({ where: { email } });
  if (!user || !(await bcrypt.compare(password, user.passwordHash)))
    throw new ApiError(401, 'Invalid identity credentials');
  res.json({ accessToken: signAccessToken({ sub: user.id, role: user.role }) });
}"""),
    'Script de despliegue': ('scripts/SOFIA-INICIAR.sh', """\
#!/bin/bash
set -e
echo "Sofia Solutions — Iniciando stack..."
APP_MODE=${APP_MODE:-secure}
docker compose down --remove-orphans
docker compose up -d --build
echo ""
echo "Stack activo:"
echo "  Frontend  -> http://localhost:8000"
echo "  Grafana   -> http://localhost:3000"
echo "  n8n       -> http://localhost:5678"
docker compose ps"""),
    'Panel SOC': ('assets/soc-metrics.js', """\
// Actualización automática de métricas SOC cada 8 segundos
async function refreshSOCMetrics() {
  const token = sessionStorage.getItem('accessToken_ADMIN');
  const res   = await fetch('/api/admin/metrics/summary', {
    headers: { 'Authorization': `Bearer ${token}` }
  });
  const data  = await res.json();
  document.getElementById('total-events').textContent    = data.totalEvents;
  document.getElementById('blocked-attacks').textContent = data.blockedAttacks;
  document.getElementById('active-sessions').textContent = data.activeSessions;
  updateChart(data.timeline);
}
setInterval(refreshSOCMetrics, 8000);
refreshSOCMetrics();"""),
}

# ── FIGURA MAP: texto del placeholder → imagen ────────────────────────────
FIGURAS = {
    'Figura 1': f'{IMG}/01_home_1778273797340.png',
    'Figura 2': f'{IMG}/11_soc_monitor_1778273920200.png',
    'Figura 3': f'{IMG}/16_client_dashboard_1778274002958.png',
    'Figura 4': f'{IMG}/11_soc_monitor_1778273920200.png',
    'Figura 5': f'{IMG}/03_admin_dashboard_1778273812760.png',
    'Figura 6': f'{IMG}/13_grafana_dashboard_1778273945380.png',
    'Figura 7': f'{IMG}/13_grafana_dashboard_final_1778273964638.png',
    'Figura 8': f'{IMG}/05_sqli_bypass_v1_success_1778273867248.png',
    'Figura 9': f'{IMG}/07_brute_force_1778273883180.png',
}

# ── Procesar párrafos ─────────────────────────────────────────────────────
paras = list(doc.paragraphs)
for i, p in enumerate(paras):
    try:
        sn = p.style.name or ''
    except:
        sn = ''
    txt = p.text.strip()

    # Insertar bloques VSCode después de Heading 3
    if sn == 'Heading 3':
        for key, (fname, code) in CODES.items():
            if key.lower() in txt.lower():
                vscode_block(p, fname, code)
                break

    # Reemplazar placeholders (Figura X: Insertar aquí...)
    if txt.startswith('(Figura') and 'Insertar' in txt:
        for fig_key, img_path in FIGURAS.items():
            if fig_key in txt:
                insert_image_after(p, img_path, width_cm=13)
                break

# ── Mejorar tablas existentes ─────────────────────────────────────────────
for table in doc.tables:
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            fill = '0E4D92' if r_idx == 0 else ('EBF2FA' if r_idx % 2 == 0 else 'FFFFFF')
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), fill); shd.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shd)
            if r_idx == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                        run.bold = True

doc.save(OUT)
print(f'OK — {OUT}')
