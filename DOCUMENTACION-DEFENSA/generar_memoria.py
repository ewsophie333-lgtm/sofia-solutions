"""
Generador de Memoria TFG - Sofia Solutions
IES Enric Soler i Godes | ASIR 2025-2026
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Helpers ────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Verdana'
        run.font.size = Pt(14 if level == 1 else 12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)
    return p

def add_body(doc, text, first_line_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.space_after = Pt(6)
    fmt.line_spacing = Pt(18)  # 1.5 aprox
    if first_line_indent:
        fmt.first_line_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Verdana'
    run.font.size = Pt(12)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = p.paragraph_format
    fmt.left_indent = Cm(1 + level * 0.5)
    fmt.line_spacing = Pt(18)
    run = p.add_run(text)
    run.font.name = 'Verdana'
    run.font.size = Pt(12)
    return p

def add_code_block(doc, filename, code):
    # Fila de encabezado con nombre de archivo
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Normal Table'
    # Header
    hdr = tbl.rows[0].cells[0]
    hdr.paragraphs[0].clear()
    hdr_run = hdr.paragraphs[0].add_run(filename)
    hdr_run.bold = True
    hdr_run.font.name = 'Verdana'
    hdr_run.font.size = Pt(10)
    hdr.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Shading header
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'D9E1F2')
    hdr._tc.get_or_add_tcPr().append(shading)
    # Code
    code_cell = tbl.rows[1].cells[0]
    code_cell.paragraphs[0].clear()
    code_run = code_cell.paragraphs[0].add_run(code)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9)
    code_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

def add_table_row(tbl, *cells, header=False):
    row = tbl.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(str(text))
        run.font.name = 'Verdana'
        run.font.size = Pt(10)
        run.bold = header
        if header:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '1F3564')
            cell._tc.get_or_add_tcPr().append(shading)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return row

def add_figure_caption(doc, n, desc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Figura {n}: {desc}')
    run.font.name = 'Verdana'
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def add_page_break(doc):
    doc.add_page_break()

def add_mermaid(doc, title, code):
    add_body(doc, f'[DIAGRAMA: {title}]', first_line_indent=False)
    add_code_block(doc, f'Mermaid — {title}', code)

# ─── Documento ──────────────────────────────────────────────────────────────

doc = Document(r'C:\Users\shair\Downloads\Memoria. Entrega 4.docx')

# Configurar márgenes en la primera sección
section = doc.sections[0]
section.left_margin   = Cm(3)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

# ─── Encabezado y pie ───────────────────────────────────────────────────────
from docx.oxml.ns import nsmap
header = section.header
if not header.paragraphs:
    header.add_paragraph()
hp = header.paragraphs[0]
hp.clear()
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run('Sofia Solutions — Laura Sofia Gomez Salazar')
hr.font.name = 'Verdana'
hr.font.size = Pt(9)
hr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

footer = section.footer
if not footer.paragraphs:
    footer.add_paragraph()
fp = footer.paragraphs[0]
fp.clear()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Número de página automático
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.text = 'PAGE'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run_pg = fp.add_run()
run_pg._r.append(fldChar1)
run_pg._r.append(instrText)
run_pg._r.append(fldChar2)
run_pg.font.name = 'Verdana'
run_pg.font.size = Pt(10)

# ─── Borrar párrafos vacíos del final y añadir contenido nuevo ──────────────
# Encontrar el índice del párrafo "Referencias" para insertar antes de él
# En su lugar, añadimos las secciones que faltan ANTES de las referencias.

# Limpiamos los párrafos vacíos al final del documento (después de las conclusiones)
# y añadimos contenido a partir del punto 7 y 8 que ya tiene contenido, así que
# solo completamos las secciones que según el índice están incompletas.

# Añadimos al final del documento la sección de soporte físico y anexos
add_page_break(doc)

# ────────────────────────────────────────────────────────────────────────────
# 9. PFC EN SUPORT FÍSIC
# ────────────────────────────────────────────────────────────────────────────
add_heading(doc, '9. PFC en Suport Físic (Pendrive USB)', level=1)
add_body(doc,
    'Junto a la memoria impresa se entrega un pendrive USB con el contenido íntegro del proyecto. '
    'La estructura de carpetas es la siguiente:')

usb_content = [
    ('Carpeta / Archivo', 'Descripción'),
    ('/sofia-solutions/', 'Directorio raíz del proyecto completo'),
    ('  /sofia-backend/', 'Código fuente del servidor Node.js + TypeScript'),
    ('  /frontend-php/', 'Interfaz web PHP (portal corporativo y SOC)'),
    ('  /frontend/', 'Interfaz React (panel de cliente)'),
    ('  /scripts/', 'Scripts de auditoría y automatización (bash)'),
    ('  /DOCUMENTACION-DEFENSA/', 'Guiones de ataque, arquitectura y métricas'),
    ('  docker-compose.yml', 'Orquestación de todos los servicios'),
    ('  .devcontainer/', 'Configuración GitHub Codespaces'),
    ('  n8n_workflow_sofia.json', 'Flujo de trabajo de alertas SOS (n8n)'),
    ('/MEMORIA_TFG.docx', 'Documento de memoria del proyecto'),
    ('/PRESENTACION_SOFIA.html', 'Presentación interactiva del proyecto'),
    ('/README.md', 'Instrucciones de despliegue y uso'),
]

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Normal Table'
add_table_row(tbl, 'Carpeta / Archivo', 'Descripción', header=True)
for row_data in usb_content[1:]:
    add_table_row(tbl, *row_data)
doc.add_paragraph()

add_body(doc,
    'Para ejecutar el proyecto completo desde el pendrive, solo es necesario tener Docker Desktop '
    'instalado y ejecutar el comando docker compose up -d desde el directorio raíz. El sistema '
    'levantará automáticamente todos los servicios: backend, frontend, base de datos PostgreSQL, '
    'panel de telemetría Grafana y el motor de automatización n8n.')

# ────────────────────────────────────────────────────────────────────────────
# ANNEX I: PLA D'EMPRESA
# ────────────────────────────────────────────────────────────────────────────
add_page_break(doc)
add_heading(doc, 'Annex I: Pla d\'Empresa', level=1)

add_body(doc,
    'Sofia Solutions nace como respuesta a una necesidad real del mercado: las pequeñas y medianas '
    'empresas carecen de los recursos económicos para contratar un equipo de ciberseguridad dedicado, '
    'pero están igualmente expuestas a las mismas amenazas que las grandes corporaciones. En este '
    'contexto, el modelo de negocio de Sofia Solutions se basa en ofrecer servicios de seguridad '
    'gestionada (MSSP) con un coste mensual adaptado al tamaño de la organización.')

add_heading(doc, 'Modelo de Negocio', level=2)
add_body(doc,
    'El modelo de negocio es de tipo SaaS (Software as a Service) con facturación mensual por niveles '
    'de servicio. Cada cliente contrata uno de los tres planes disponibles según sus necesidades:')

planes = [
    ('Plan', 'Precio/mes', 'Servicios incluidos'),
    ('Essential', '299 €/mes', 'Monitorización 8x5, alertas email, 1 auditoría trimestral'),
    ('Business', '699 €/mes', 'Monitorización 24x7, alertas SMS/email, 2 auditorías/año, SOC dashboard'),
    ('Enterprise', '1.499 €/mes', 'SOC dedicado 24x7, IR Retainer, auditorías ilimitadas, SLA 15 min'),
]

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Normal Table'
add_table_row(tbl2, 'Plan', 'Precio/mes', 'Servicios incluidos', header=True)
for row_data in planes[1:]:
    add_table_row(tbl2, *row_data)
doc.add_paragraph()

add_heading(doc, 'Mercado Objetivo', level=2)
add_body(doc,
    'El público objetivo de Sofia Solutions son PYMEs de entre 10 y 200 empleados en sectores '
    'regulados como finanzas, seguros, sanidad e industria crítica. Estos sectores están sujetos '
    'a normativas como ISO 27001, ENS o GDPR que obligan a implementar controles de seguridad '
    'auditables. Sofia Solutions facilita el cumplimiento normativo a un coste accesible.')

add_heading(doc, 'Análisis DAFO', level=2)

dafo = [
    ('Debilidades', 'Amenazas'),
    ('• Equipo pequeño (1 persona en fase inicial)', '• Competencia con MSSPs establecidos'),
    ('• Dependencia de herramientas open source', '• Cambios normativos frecuentes'),
    ('Fortalezas', 'Oportunidades'),
    ('• Bajo coste de infraestructura (Docker/Open Source)', '• Crecimiento exponencial de ciberataques a PYMEs'),
    ('• Plataforma 100% desarrollada internamente', '• Demanda de servicios SOC asequibles en España'),
    ('• Despliegue cloud-native sin hardware propio', '• Subvenciones europeas para digitalización (NextGenEU)'),
]

tbl3 = doc.add_table(rows=4, cols=2)
tbl3.style = 'Normal Table'
labels = ['Debilidades', 'Amenazas', 'Fortalezas', 'Oportunidades']
content = [
    '• Equipo pequeño (1 persona en fase inicial)\n• Dependencia de herramientas open source',
    '• Competencia con MSSPs establecidos\n• Cambios normativos frecuentes',
    '• Bajo coste de infraestructura (Docker/Open Source)\n• Plataforma 100% desarrollada internamente',
    '• Crecimiento exponencial de ciberataques a PYMEs\n• Demanda de servicios SOC asequibles en España',
]
for i, row in enumerate(tbl3.rows):
    col = i % 2
    row_idx = i // 2
    for j, cell in enumerate(row.cells):
        cell.paragraphs[0].clear()
        idx = i * 2 + j
        if idx < len(labels):
            label_run = cell.paragraphs[0].add_run(labels[idx] + '\n')
            label_run.bold = True
            label_run.font.name = 'Verdana'
            label_run.font.size = Pt(10)
            content_run = cell.paragraphs[0].add_run(content[idx])
            content_run.font.name = 'Verdana'
            content_run.font.size = Pt(10)

doc.add_paragraph()

add_heading(doc, 'Proyección Económica (Año 1)', level=2)

proyeccion = [
    ('Concepto', 'Importe anual'),
    ('Ingresos estimados (5 clientes Essential + 2 Business)', '26.580 €'),
    ('Costes de infraestructura (VPS, dominio, certificados)', '1.200 €'),
    ('Costes de herramientas (n8n Cloud, Grafana Cloud)', '960 €'),
    ('Costes de marketing y presencia web', '600 €'),
    ('Margen bruto estimado', '23.820 €'),
]

tbl4 = doc.add_table(rows=1, cols=2)
tbl4.style = 'Normal Table'
add_table_row(tbl4, 'Concepto', 'Importe anual', header=True)
for row_data in proyeccion[1:]:
    add_table_row(tbl4, *row_data)
doc.add_paragraph()

add_body(doc,
    'Estos datos son una estimación conservadora basada en el coste real de los servicios cloud '
    'utilizados durante el desarrollo del proyecto. La viabilidad económica del modelo es clara '
    'dado que la inversión inicial es prácticamente nula al usar herramientas de código abierto '
    'y contenedores Docker como infraestructura de despliegue.')

# ─── Guardar ────────────────────────────────────────────────────────────────
output_path = r'C:\Users\shair\Downloads\Memoria_Sofia_COMPLETA.docx'
doc.save(output_path)
print(f'OK - Documento guardado en: {output_path}')
